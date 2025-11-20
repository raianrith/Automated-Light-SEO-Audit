import sys
import os
import platform

# Fix architecture conflicts by prioritizing local and user packages (macOS only)
# This code is only needed on macOS with architecture conflicts, skip on Streamlit Cloud (Linux)
if platform.system() == 'Darwin':  # macOS only
    # Remove broken system-wide site-packages from path
    system_site_packages = '/Library/Frameworks/Python.framework/Versions/3.13/lib/python3.13/site-packages'
    if system_site_packages in sys.path:
        sys.path.remove(system_site_packages)

    # Add local packages directory first (contains arm64-compatible Pillow)
    local_packages = os.path.join(os.path.dirname(__file__), '.local_packages')
    if os.path.exists(local_packages):
        sys.path.insert(0, local_packages)

    # Then prioritize user site-packages
    user_site = None
    for path in sys.path:
        if 'Users' in path and 'site-packages' in path:
            user_site = path
            break
    if user_site and user_site not in sys.path[:2]:
        sys.path.insert(0, user_site)

    # CRITICAL: Import PIL BEFORE matplotlib tries to import it
    # This ensures matplotlib uses the correct arm64-compatible version from user site-packages
    # Remove any broken PIL from sys.modules
    for key in list(sys.modules.keys()):
        if key.startswith('PIL'):
            del sys.modules[key]

    # Import PIL now (will use user site-packages since system is removed from path)
    try:
        from PIL import Image
        # Cache it so matplotlib finds it
        sys.modules['PIL'] = __import__('PIL')
        sys.modules['PIL.Image'] = Image
    except ImportError:
        # If that fails, try with local_packages
        if os.path.exists(local_packages) and local_packages not in sys.path:
            sys.path.insert(0, local_packages)
            try:
                from PIL import Image
                sys.modules['PIL'] = __import__('PIL')
                sys.modules['PIL.Image'] = Image
            except ImportError:
                pass

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import numpy as np
import re
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from openai import OpenAI


# Configure page
st.set_page_config(
    page_title="SEO Performance Analyzer",
    page_icon="📈",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2rem;
        font-weight: 600;
        color: #2c3e50;
        text-align: center;
        margin-bottom: 2rem;
        padding: 1rem 0;
        background-color: #f9f9f9;
        border-radius: 12px;
        box-shadow: 0 2px 6px rgba(0,0,0,0.08);
        letter-spacing: 0.5px;
    }
    .main-header span {
        color: #1f77b4; /* accent color */
    }
    
    .section-header {
        font-size: 1.6rem;
        font-weight: 600;
        color: #2c3e50;
        margin: 2rem 0 1rem 0;
        padding-bottom: 0.3rem;
        border-bottom: 2px solid #3498db; /* subtle accent */
        letter-spacing: 0.5px;
    }

        
    .instruction-box {
        background-color: #f0f8ff;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #2196f3;
        margin: 1rem 0;
    }
    
    .insight-box {
        background-color: #f8fff8;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #4caf50;
        margin: 1rem 0;
    }
    
    .warning-box {
        background-color: #fff3cd;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #ffc107;
        margin: 1rem 0;
    }
    
    .metric-card {
        background-color: white;
        padding: 1rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        border: 1px solid #e0e0e0;
    }
    
    .file-upload-section {
        background-color: #fafafa;
        padding: 1.5rem;
        border-radius: 10px;
        border: 2px dashed #ccc;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

def main():
    # Main header
    st.markdown('<div class="main-header">🚀 <span>SEO Performance Analyzer</span></div>', unsafe_allow_html=True)
    
    # Sidebar guide
    with st.sidebar:
        st.markdown("### 📚 How It Works")
        st.markdown("""
        1. **Choose** an analysis tab above
        2. **Upload** required CSV/Excel files  
        3. **Review** automated insights
        4. **Download** your reports
        """)
        
        st.markdown("---")
        
        st.markdown("### 🔧 Data Sources")
        st.markdown("""
        - **Semrush**: Keyword rankings & competition data
        - **Google Search Console**: Click & impression metrics
        - **GA4**: Traffic & conversion analytics
        """)
        
        st.markdown("---")
        
        st.markdown("### 💡 Pro Tips")
        st.markdown("""
        - Export in **CSV or Excel format** (never PDF)
        - Use **consistent date ranges** across all exports
        - **Same month comparisons** for YoY analysis
        - Check column headers match expectations
        """)
    
    # Enhanced tab navigation with more sections
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "📋 Data Export Guide",
        "📊 Visibility Trends", 
        "🔄 Keyword Movement", 
        "📄 Page Performance",
        "🎯 Query Analysis",
        "🏁 Competitor Gaps"
    ])
    
    with tab1:
        data_export_instructions()
        
    with tab2:
        keyword_visibility_analysis()
        
    with tab3:
        keyword_movement_analysis()
        
    with tab4:
        page_performance_analysis()
        
    with tab5:
        query_gains_losses_analysis()
        
    with tab6:
        competitor_analysis()
    
# Helper functions for file processing

# ===================== REPORT HELPERS & BUILDER (PUT ABOVE comprehensive_report_tab) =====================

@st.cache_data(show_spinner=False)
def _load_df_cached(file_bytes, file_name, file_type):
    """Cached file reading - converts uploaded file to bytes for caching"""
    try:
        if file_type == "csv":
            df = pd.read_csv(io.BytesIO(file_bytes))
        elif file_type in ["xlsx", "xls"]:
            df = pd.read_excel(io.BytesIO(file_bytes))
        else:
            df = pd.read_csv(io.BytesIO(file_bytes))
        # normalize columns
        df.columns = [str(c).strip().lower().replace(" ", "_") for c in df.columns]
        return df
    except Exception as e:
        st.error(f"Error reading file {file_name}: {str(e)}")
        return None

def _load_df(uploaded_file):
    """Read CSV/XLS/XLSX to DataFrame with normalized columns. Returns None if not provided/failed."""
    if not uploaded_file:
        return None
    name = getattr(uploaded_file, "name", "").lower()
    try:
        # Read file into memory for caching
        uploaded_file.seek(0)
        file_bytes = uploaded_file.read()
        
        # Determine file type
        if name.endswith(".csv"):
            file_type = "csv"
        elif name.endswith(".xlsx"):
            file_type = "xlsx"
        elif name.endswith(".xls"):
            file_type = "xls"
        else:
            file_type = "csv"  # fallback
        
        # Use cached version
        return _load_df_cached(file_bytes, name, file_type)
    except Exception as e:
        st.error(f"Error reading file {name}: {str(e)}")
        return None


# ---------- KPI formatting ----------
def _k(v, default="—"):
    if v is None:
        return default
    try:
        if isinstance(v, float):
            if abs(v) >= 1000:
                return f"{v:,.0f}"
            return f"{v:,.2f}"
        if isinstance(v, (int, np.integer)):
            return f"{int(v):,}"
        return str(v)
    except Exception:
        return str(v)


# ---------- Matplotlib -> PNG bytes ----------
def _save_current_fig(width_in=6, height_in=3.2, dpi=140):
    buf = io.BytesIO()
    fig = plt.gcf()
    fig.set_size_inches(width_in, height_in)
    plt.tight_layout()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()

def _bar(labels, vals, title=None, xlabel=None, ylabel=None, rotate=0):
    plt.figure()
    plt.bar(labels, vals)
    if rotate:
        plt.xticks(rotation=rotate, ha="right")
    if title: plt.title(title)
    if xlabel: plt.xlabel(xlabel)
    if ylabel: plt.ylabel(ylabel)
    return _save_current_fig()

def _barh(labels, vals, title=None, xlabel=None, ylabel=None):
    plt.figure()
    plt.barh(labels, vals)
    if title: plt.title(title)
    if xlabel: plt.xlabel(xlabel)
    if ylabel: plt.ylabel(ylabel)
    return _save_current_fig()

def _pie(labels, vals, title=None):
    plt.figure()
    plt.pie(vals, labels=labels, autopct="%1.0f%%", startangle=90)
    plt.axis("equal")
    if title: plt.title(title)
    return _save_current_fig()


# ---------- docx helpers ----------
def _docx_title(doc: Document, title: str, subtitle: str | None = None):
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _docx_kpi_row(doc: Document, kpis: dict):
    table = doc.add_table(rows=1, cols=len(kpis))
    table.style = "Light List"
    row = table.rows[0].cells
    for i, (k, v) in enumerate(kpis.items()):
        row[i].text = f"{k}\n{_k(v)}"

def _docx_add_image(doc: Document, image_bytes: bytes, caption: str | None = None, width_in=6):
    if not image_bytes:
        return
    doc.add_picture(io.BytesIO(image_bytes), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _docx_note(doc: Document, text: str, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic


# ---------- sections (only the 7 sources you picked) ----------
def _section_semrush_visibility(doc: Document, current_df: pd.DataFrame | None, prev_df: pd.DataFrame | None):
    if current_df is None and prev_df is None:
        return
    doc.add_heading("Overall Keyword Visibility (YoY)", level=1)

    def top_share(df, cutoff):
        if df is None or df.empty:
            return None
        pos_col = next((c for c in df.columns if c in {"position", "pos", "rank"}), None)
        if not pos_col:  # try fuzzy match
            pos_col = next((c for c in df.columns if "pos" in c), None)
        if not pos_col:
            return None
        total = len(df)
        top = (pd.to_numeric(df[pos_col], errors="coerce") <= cutoff).sum()
        return (top / total) * 100 if total > 0 else None

    kpis = {
        "Keywords (Current)": len(current_df) if current_df is not None else None,
        "Top 3 % (Current)": top_share(current_df, 3),
        "Top 10 % (Current)": top_share(current_df, 10),
        "Keywords (Prev)": len(prev_df) if prev_df is not None else None,
        "Top 3 % (Prev)": top_share(prev_df, 3),
        "Top 10 % (Prev)": top_share(prev_df, 10),
    }
    _docx_kpi_row(doc, kpis)

    if current_df is not None and prev_df is not None:
        labels = ["Current Keywords", "Previous Keywords"]
        vals = [len(current_df), len(prev_df)]
        try:
            img = _bar(labels, vals, title="Keyword Footprint", ylabel="Count")
            _docx_add_image(doc, img, caption="Total Keywords (Current vs Previous)")
        except Exception:
            pass

    _docx_note(doc, "Interpretation: Compare footprint and quality YoY (higher Top 3/Top 10 shares are better).", italic=True)


def _section_semrush_winners_losers(doc: Document, changes_df: pd.DataFrame | None, top_n=10):
    if changes_df is None or changes_df.empty:
        return
    doc.add_heading("Top Gainers & Decliners (Semrush Position Changes)", level=1)

    df = changes_df.copy()
    kw_col = next((c for c in df.columns if "keyword" in c), None)
    from_col = next((c for c in df.columns if "from" in c and "pos" in c), None)
    to_col = next((c for c in df.columns if "to" in c and "pos" in c), None)

    if from_col and to_col:
        df["delta_pos"] = pd.to_numeric(df[from_col], errors="coerce") - pd.to_numeric(df[to_col], errors="coerce")
    else:
        df["delta_pos"] = np.nan

    winners = df.sort_values("delta_pos", ascending=False).head(top_n)
    losers = df.sort_values("delta_pos", ascending=True).head(top_n)

    # winners
    if not winners.empty and "delta_pos" in winners:
        labels = list((winners[kw_col] if kw_col else winners.index).astype(str))
        vals = list(pd.to_numeric(winners["delta_pos"], errors="coerce").fillna(0).values)
        img = _barh(labels[::-1], vals[::-1], title="Top Winners (positions improved)", xlabel="Positions")
        _docx_add_image(doc, img)

    # losers
    if not losers.empty and "delta_pos" in losers:
        labels = list((losers[kw_col] if kw_col else losers.index).astype(str))
        vals = list((-pd.to_numeric(losers["delta_pos"], errors="coerce")).fillna(0).values)
        img = _barh(labels[::-1], vals[::-1], title="Top Losers (positions dropped)", xlabel="Positions")
        _docx_add_image(doc, img)

    _docx_note(doc, "Interpretation: Gains on transactional terms vs losses on informational terms can reflect SERP/intent shifts.", italic=True)


def _section_semrush_pages(doc: Document, pages_df: pd.DataFrame | None, top_n=10):
    if pages_df is None or pages_df.empty:
        return
    doc.add_heading("Top Pages by Estimated Visits (Semrush Pages)", level=1)
    url_col = next((c for c in pages_df.columns if "url" in c or "page" in c), None)
    visits_col = next((c for c in pages_df.columns if "visit" in c or "traffic" in c or "est" in c), None)
    if not url_col or not visits_col:
        _docx_note(doc, "Skipped: required columns not found (URL/Visits).", italic=True)
        return
    top = pages_df.sort_values(visits_col, ascending=False).head(top_n)
    labels = list(top[url_col].astype(str).values)
    vals = list(pd.to_numeric(top[visits_col], errors="coerce").fillna(0).values)
    img = _barh(labels[::-1], vals[::-1], title="Top Pages (Estimated Visits)", xlabel="Visits")
    _docx_add_image(doc, img)
    _docx_note(doc, "Why it matters: Protect & grow commercial-intent pages; support with internal links and content.", italic=True)


def _section_semrush_competitors(doc: Document, competitors_df: pd.DataFrame | None):
    if competitors_df is None or competitors_df.empty:
        return
    doc.add_heading("Competitor Benchmark (Semrush)", level=1)
    name_col = next((c for c in competitors_df.columns if "domain" in c or "competitor" in c), None)
    traffic_col = next((c for c in competitors_df.columns if "traffic" in c or "visit" in c), None)

    if name_col and traffic_col:
        tbl = competitors_df[[name_col, traffic_col]].dropna().head(8)
        labels = list(tbl[name_col].astype(str).values)
        vals = list(pd.to_numeric(tbl[traffic_col], errors="coerce").fillna(0).values)
        img = _barh(labels[::-1], vals[::-1], title="Competitor Estimated Traffic", xlabel="Visits")
        _docx_add_image(doc, img)
    _docx_note(doc, "Interpretation: Higher traffic with similar overlap suggests authority/content depth gaps.", italic=True)


def _section_gsc_queries(doc: Document, queries_df: pd.DataFrame | None):
    if queries_df is None or queries_df.empty:
        return
    doc.add_heading("GSC Query Trends", level=1)
    click_col = next((c for c in queries_df.columns if "click" in c), None)
    imp_col = next((c for c in queries_df.columns if "impression" in c), None)
    query_col = next((c for c in queries_df.columns if "query" in c), None)

    kpis = {}
    if click_col: kpis["Clicks"] = pd.to_numeric(queries_df[click_col], errors="coerce").sum()
    if imp_col:   kpis["Impressions"] = pd.to_numeric(queries_df[imp_col], errors="coerce").sum()
    if kpis: _docx_kpi_row(doc, kpis)

    if click_col and query_col:
        top = queries_df.sort_values(click_col, ascending=False).head(10)
        labels = list(top[query_col].astype(str).values)
        vals = list(pd.to_numeric(top[click_col], errors="coerce").fillna(0).values)
        img = _barh(labels[::-1], vals[::-1], title="Top Queries by Clicks", xlabel="Clicks")
        _docx_add_image(doc, img)

    _docx_note(doc, "Interpretation: Impressions up with clicks down → CTR erosion from richer SERPs or rank shifts.", italic=True)


def _section_gsc_pages(doc: Document, pages_df: pd.DataFrame | None):
    if pages_df is None or pages_df.empty:
        return
    doc.add_heading("GSC Pages (Top Performers / Slipping)", level=1)
    url_col = next((c for c in pages_df.columns if "page" in c or "url" in c), None)
    click_col = next((c for c in pages_df.columns if "click" in c), None)

    if url_col and click_col:
        top = pages_df.sort_values(click_col, ascending=False).head(10)
        labels = list(top[url_col].astype(str).values)
        vals = list(pd.to_numeric(top[click_col], errors="coerce").fillna(0).values)
        img = _barh(labels[::-1], vals[::-1], title="Top Pages by Clicks", xlabel="Clicks")
        _docx_add_image(doc, img)
        _docx_note(doc, "Why it matters: Optimize top landers for conversion paths (CTAs, forms, speed).", italic=True)
    else:
        _docx_note(doc, "Skipped: required columns not found (URL/Clicks).", italic=True)


def build_seo_audit_docx(
    site_domain: str,
    semrush_current_df: pd.DataFrame | None = None,
    semrush_prev_df: pd.DataFrame | None = None,
    semrush_changes_df: pd.DataFrame | None = None,
    semrush_pages_df: pd.DataFrame | None = None,
    semrush_comp_df: pd.DataFrame | None = None,
    gsc_queries_df: pd.DataFrame | None = None,
    gsc_pages_df: pd.DataFrame | None = None,
) -> bytes:
    """Master builder that composes all sections. Skips any missing dataset gracefully."""
    doc = Document()
    _docx_title(doc, f"SEO Performance Analysis for {site_domain}",
                subtitle=datetime.now().strftime("%Y-%m-%d"))

    doc.add_heading("Executive Summary", level=1)
    _docx_note(doc,
        f"{site_domain} SEO performance summary based on uploaded Semrush and GSC exports. "
        "Sections without data were skipped. Charts and KPIs follow for each available source."
    )

    _section_semrush_visibility(doc, semrush_current_df, semrush_prev_df)
    _section_semrush_winners_losers(doc, semrush_changes_df)
    _section_semrush_pages(doc, semrush_pages_df)
    _section_semrush_competitors(doc, semrush_comp_df)
    _section_gsc_queries(doc, gsc_queries_df)
    _section_gsc_pages(doc, gsc_pages_df)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
# ===================== END REPORT HELPERS & BUILDER =====================

def read_uploaded_file(uploaded_file):
    """Read uploaded CSV or Excel file - uses cached version"""
    if uploaded_file is not None:
        return _load_df(uploaded_file)
    return None

def normalize_columns(df):
    """Normalize column names by cleaning whitespace and special characters"""
    df = df.copy()
    df.columns = [re.sub(r"\s+", " ", str(c).replace("\xa0", " ")).strip() for c in df.columns]
    return df

def find_column(columns, patterns):
    """Find column by searching for patterns (case-insensitive, handles normalized columns)"""
    # Handle both normalized (lowercase with underscores) and original column names
    columns_lower = {str(c).lower().replace(" ", "_").replace("-", "_"): c for c in columns}
    columns_normalized = {str(c).lower().replace(" ", "_").replace("-", "_"): c for c in columns}
    
    # Combine both mappings
    all_mappings = {**columns_lower, **columns_normalized}
    
    for pattern in patterns:
        pattern_normalized = pattern.lower().replace(" ", "_").replace("-", "_")
        # Try exact match first
        if pattern_normalized in all_mappings:
            return all_mappings[pattern_normalized]
        # Try substring match
        for col_normalized, original_col in all_mappings.items():
            if pattern_normalized in col_normalized or col_normalized in pattern_normalized:
                return original_col
    return None

def preview_columns(df, file_name="File"):
    """Preview detected columns for user validation"""
    if df is None or len(df.columns) == 0:
        return None
    
    with st.expander(f"📋 Column Detection Preview - {file_name}", expanded=False):
        st.markdown(f"**Detected {len(df.columns)} columns:**")
        cols_display = pd.DataFrame({
            'Column Name': df.columns,
            'Sample Value': [str(df[col].iloc[0])[:50] if len(df) > 0 and pd.notna(df[col].iloc[0]) else 'N/A' for col in df.columns]
        })
        st.dataframe(cols_display, use_container_width=True, hide_index=True, height=200)
        
        # Show detected key columns
        kw_col = find_column(df.columns, ['keyword', 'keywords', 'query', 'queries'])
        pos_col = find_column(df.columns, ['position', 'pos', 'rank', 'ranking'])
        if kw_col or pos_col:
            st.success(f"✅ Key columns detected: Keyword={kw_col or 'Not found'}, Position={pos_col or 'Not found'}")
        
        return cols_display

def ensure_standardized_columns(df):
    """Ensure DataFrame has standardized column names for consistent processing across clients"""
    if df is None:
        return df
    
    df = df.copy()
    # Normalize to lowercase with underscores (consistent with _load_df_cached)
    df.columns = [str(c).strip().lower().replace(" ", "_").replace("-", "_") for c in df.columns]
    return df

def keyword_visibility_analysis():
    st.markdown('<div class="section-header">🔍 Keyword Visibility Trends (Year-over-Year)</div>', unsafe_allow_html=True)
    
    # Modern instruction design using containers and columns
    with st.container():
        st.markdown("### 📊 Analysis Overview")
        
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown("""
            This analysis compares your keyword rankings between two time periods to understand:
            
            **🎯 Key Questions Answered:**
            - Are you ranking for more or fewer keywords?
            - What percentage of keywords are in top positions? 
            - Are you gaining authority or losing visibility breadth?
            """)
        
        with col2:
            st.info("""
            **💡 Strategic Value**
            
            Identifies whether you're building stronger authority or need to expand your keyword footprint.
            """)
    
    # File requirements in expandable section
    with st.expander("📁 **File Requirements & Setup**", expanded=False):
        st.markdown("""
        **Required Files:** 2 Semrush Positions exports
        
        | File | Description | Export From |
        |------|-------------|-------------|
        | **Current Period** | Recent month Semrush Positions | Domain Analytics → Organic Research → Positions |
        | **Previous Period** | Same month last year | Same location, different date |
        
        **📋 Export Settings:**
        - Database: United States (or your target country)
        - Device: Desktop 
        - Format: CSV or Excel
        - Date: Current month vs Same month last year
        """)
    
    # Key insights preview
    st.markdown("### 🎯 Analysis Insights You'll Get")
    
    insight_col1, insight_col2, insight_col3 = st.columns(3)
    
    with insight_col1:
        st.markdown("""
        **📈 Total Keywords**
        - Year-over-year change
        - Growth vs decline analysis
        """)
    
    with insight_col2:
        st.markdown("""
        **🏆 Ranking Quality**
        - Top 3, 4-10, 11-20, 21+ distribution
        - Quality vs quantity trade-offs
        """)
    
    with insight_col3:
        st.markdown("""
        **💡 Strategic Recommendations**
        - Authority building opportunities
        - Breadth expansion needs
        """)
    
    st.markdown("---")
    
    # File upload section
    st.markdown("### 📤 Upload Your Data Files")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 📊 Current Period (2024/2025)")
        current_file = st.file_uploader(
            "Upload current Semrush Positions file",
            type=['csv', 'xlsx', 'xls'],
            key="current_positions",
            help="Export from Semrush: Domain Analytics → Organic Research → Positions (CSV or Excel format)"
        )
        
    with col2:
        st.markdown("#### 📊 Previous Period (Same Month Last Year)")
        previous_file = st.file_uploader(
            "Upload previous year Semrush Positions file", 
            type=['csv', 'xlsx', 'xls'],
            key="previous_positions",
            help="Same export but for the corresponding month last year (CSV or Excel format)"
        )
    
    # Process files if both are uploaded
    if current_file is not None and previous_file is not None:
        # Add Run Analysis button (centered)
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            run_analysis = st.button("🚀 Run Visibility Analysis", key="run_visibility", type="primary", use_container_width=True)
        
        # Display results outside column context for full width
        if run_analysis:
            with st.spinner("🔄 Processing your data..."):
                try:
                    # Load data using helper functions
                    current_df = ensure_standardized_columns(read_uploaded_file(current_file))
                    previous_df = ensure_standardized_columns(read_uploaded_file(previous_file))
                    
                    # Show column preview for user validation
                    if current_df is not None:
                        preview_columns(current_df, current_file.name if hasattr(current_file, 'name') else "Current File")
                    if previous_df is not None:
                        preview_columns(previous_df, previous_file.name if hasattr(previous_file, 'name') else "Previous File")
                    
                    # Validate data
                    validation_passed, validation_message = validate_positions_data(current_df, previous_df)
                    
                    if not validation_passed:
                        st.error(validation_message)
                        st.info("💡 **Tip:** Column names are automatically detected. If detection fails, check the column preview above to see what columns were found.")
                        st.stop()
                    
                    # Perform analysis
                    analysis_results = analyze_keyword_visibility(current_df, previous_df)
                    
                    # Display results - NOW IN FULL WIDTH
                    display_visibility_results(analysis_results)
                    
                except Exception as e:
                    st.error(f"❌ Error processing files: {str(e)}")
                    st.info("💡 Please ensure you've uploaded valid Semrush Positions CSV or Excel files")
    else:
        if current_file is None:
            st.info("📤 Please upload the current period Semrush Positions file")
        if previous_file is None:
            st.info("📤 Please upload the previous period Semrush Positions file")

def validate_positions_data(current_df, previous_df):
    """Validate the uploaded Semrush positions data with flexible column detection"""
    
    # Use find_column for flexible detection (works with normalized columns)
    for df, period in [(current_df, 'current'), (previous_df, 'previous')]:
        # Check for keyword column (handles various naming)
        kw_col = find_column(df.columns, ['keyword', 'keywords', 'query', 'queries'])
        pos_col = find_column(df.columns, ['position', 'pos', 'rank', 'ranking'])
        
        missing_columns = []
        if not kw_col:
            missing_columns.append('Keyword/Query')
        if not pos_col:
            missing_columns.append('Position/Rank')
        
        if missing_columns:
            available = ", ".join(list(df.columns)[:15])
            return False, f"❌ Missing required columns in {period} file: {missing_columns}. Available columns: {available}..."
    
    # Check if data is not empty
    if len(current_df) == 0 or len(previous_df) == 0:
        return False, "❌ One or both files appear to be empty"
    
    # Check for valid position data using detected column
    for df, period in [(current_df, 'current'), (previous_df, 'previous')]:
        pos_col = find_column(df.columns, ['position', 'pos', 'rank', 'ranking'])
        if pos_col and df[pos_col].isna().all():
            return False, f"❌ No valid position data found in {period} file"
    
    return True, "✅ Data validation passed - columns detected successfully"

@st.cache_data(show_spinner=False)
def _analyze_keyword_visibility_cached(current_df_hash, previous_df_hash, current_df, previous_df):
    """Cached analysis function - uses DataFrame hash for cache key"""
    # Detect column names dynamically
    pos_col_current = find_column(current_df.columns, ['position', 'pos', 'rank', 'ranking']) or 'position'
    pos_col_prev = find_column(previous_df.columns, ['position', 'pos', 'rank', 'ranking']) or 'position'
    
    # Clean position data - convert to numeric, handle non-numeric values (vectorized)
    def clean_positions(df, pos_col):
        df = df.copy()
        # Find position column (handles normalized column names)
        if pos_col not in df.columns:
            pos_col = find_column(df.columns, ['position', 'pos', 'rank', 'ranking']) or 'position'
        # Use lowercase normalized name
        pos_col_normalized = pos_col.lower().replace(" ", "_").replace("-", "_")
        if pos_col_normalized in df.columns:
            pos_col = pos_col_normalized
        elif pos_col in df.columns:
            pass  # Use as-is
        else:
            # Try to find any position-like column
            pos_col = find_column(df.columns, ['position', 'pos', 'rank', 'ranking']) or 'position'
        
        if pos_col in df.columns:
            df['position'] = pd.to_numeric(df[pos_col], errors='coerce')
            df = df.dropna(subset=['position'])
        return df
    
    current_clean = clean_positions(current_df, pos_col_current)
    previous_clean = clean_positions(previous_df, pos_col_prev)
    
    # Calculate total keywords
    total_current = len(current_clean)
    total_previous = len(previous_clean)
    
    # Calculate rank buckets using vectorized operations
    def get_rank_buckets(df):
        # Use normalized column name
        pos_col = 'position' if 'position' in df.columns else find_column(df.columns, ['position', 'pos', 'rank', 'ranking']) or 'position'
        if pos_col not in df.columns:
            return {'top_3': 0, 'top_4_10': 0, 'top_11_20': 0, 'top_21_plus': 0}
        pos = df[pos_col]
        return {
            'top_3': (pos <= 3).sum(),
            'top_4_10': ((pos > 3) & (pos <= 10)).sum(),
            'top_11_20': ((pos > 10) & (pos <= 20)).sum(),
            'top_21_plus': (pos > 20).sum()
        }
    
    current_buckets = get_rank_buckets(current_clean)
    previous_buckets = get_rank_buckets(previous_clean)
    
    # Calculate changes
    total_change = total_current - total_previous
    total_change_pct = (total_change / total_previous * 100) if total_previous > 0 else 0
    
    # Calculate bucket changes
    bucket_changes = {}
    for bucket in current_buckets:
        current_count = current_buckets[bucket]
        previous_count = previous_buckets[bucket]
        change = current_count - previous_count
        change_pct = (change / previous_count * 100) if previous_count > 0 else 0
        
        bucket_changes[bucket] = {
            'current': current_count,
            'previous': previous_count,
            'change': change,
            'change_pct': change_pct,
            'current_share': (current_count / total_current * 100) if total_current > 0 else 0,
            'previous_share': (previous_count / total_previous * 100) if total_previous > 0 else 0
        }
    
    return {
        'total_current': total_current,
        'total_previous': total_previous,
        'total_change': total_change,
        'total_change_pct': total_change_pct,
        'bucket_changes': bucket_changes,
        'current_df': current_clean,
        'previous_df': previous_clean
    }

def analyze_keyword_visibility(current_df, previous_df):
    """Analyze keyword visibility trends - uses cached version"""
    # Create hash for caching (using shape and first few rows)
    current_hash = hash((current_df.shape, tuple(current_df.head(5).values.flatten()) if len(current_df) > 0 else ()))
    previous_hash = hash((previous_df.shape, tuple(previous_df.head(5).values.flatten()) if len(previous_df) > 0 else ()))
    
    return _analyze_keyword_visibility_cached(current_hash, previous_hash, current_df, previous_df)

def display_visibility_results(results):
    """Display the keyword visibility analysis results"""
    
    # Key metrics section
    st.markdown('<div class="section-header">📈 Key Metrics</div>', unsafe_allow_html=True)
    
    # Metrics row
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        # For keywords: negative change is bad (red), positive is good (green)
        # Streamlit's "normal" mode: positive = green, negative = red (correct for this metric)
        st.metric(
            label="Total Keywords",
            value=f"{results['total_current']:,}",
            delta=f"{results['total_change']:,} ({results['total_change_pct']:.1f}%)",
            delta_color="normal"  # Always normal: negative changes show red (bad), positive show green (good)
        )
    
    with col2:
        top_3_current = results['bucket_changes']['top_3']['current_share']
        top_3_previous = results['bucket_changes']['top_3']['previous_share']
        top_3_delta = top_3_current - top_3_previous
        # For rankings: positive change is good (green), negative is bad (red)
        # Streamlit's "normal" mode handles this correctly: positive = green, negative = red
        st.metric(
            label="Top 3 Rankings",
            value=f"{top_3_current:.1f}%",
            delta=f"{top_3_delta:+.1f}%",
            delta_color="normal"  # Always normal: positive = green (good), negative = red (bad)
        )
    
    with col3:
        top_10_current = results['bucket_changes']['top_3']['current_share'] + results['bucket_changes']['top_4_10']['current_share']
        top_10_previous = results['bucket_changes']['top_3']['previous_share'] + results['bucket_changes']['top_4_10']['previous_share']
        top_10_delta = top_10_current - top_10_previous
        # For rankings: positive change is good (green), negative is bad (red)
        # Streamlit's "normal" mode handles this correctly: positive = green, negative = red
        st.metric(
            label="Top 10 Rankings",
            value=f"{top_10_current:.1f}%",
            delta=f"{top_10_delta:+.1f}%",
            delta_color="normal"  # Always normal: positive = green (good), negative = red (bad)
        )
    
    with col4:
        # Quality score (weighted average of rankings)
        def calc_quality_score(buckets, total):
            if total == 0:
                return 0
            score = ((buckets['top_3'] * 100) + 
                    (buckets['top_4_10'] * 75) + 
                    (buckets['top_11_20'] * 50) + 
                    (buckets['top_21_plus'] * 25)) / total
            return score
        
        current_quality = calc_quality_score(
            {k: v['current'] for k, v in results['bucket_changes'].items()}, 
            results['total_current']
        )
        previous_quality = calc_quality_score(
            {k: v['previous'] for k, v in results['bucket_changes'].items()}, 
            results['total_previous']
        )
        quality_delta = current_quality - previous_quality
        # For quality score: positive change is good (green), negative is bad (red)
        # Streamlit's "normal" mode handles this correctly: positive = green, negative = red
        st.metric(
            label="Quality Score",
            value=f"{current_quality:.1f}",
            delta=f"{quality_delta:+.1f}",
            delta_color="normal",  # Always normal: positive = green (good), negative = red (bad)
            help="Weighted score: Top 3=100pts, 4-10=75pts, 11-20=50pts, 21+=25pts"
        )
    
    # Visualization section
    st.markdown('<div class="section-header">📊 Ranking Distribution Analysis</div>', unsafe_allow_html=True)
    
    # Prepare data for charts
    bucket_labels = ['Top 3', '4-10', '11-20', '21+']
    current_values = [results['bucket_changes'][k]['current'] for k in ['top_3', 'top_4_10', 'top_11_20', 'top_21_plus']]
    previous_values = [results['bucket_changes'][k]['previous'] for k in ['top_3', 'top_4_10', 'top_11_20', 'top_21_plus']]
    
    # Full width charts - stacked vertically for better space usage
    # First chart - ranking distribution comparison
    fig_distribution = go.Figure(data=[
        go.Bar(name='Previous Period', x=bucket_labels, y=previous_values, 
               marker_color='lightblue', text=[f"{val:,}" for val in previous_values],
               textposition='auto', textfont=dict(size=12)),
        go.Bar(name='Current Period', x=bucket_labels, y=current_values, 
               marker_color='darkblue', text=[f"{val:,}" for val in current_values],
               textposition='auto', textfont=dict(size=12))
    ])
    
    fig_distribution.update_layout(
        title=dict(text='Keyword Count by Ranking Position', font=dict(size=20)),
        xaxis_title='Ranking Position',
        yaxis_title='Number of Keywords',
        barmode='group',
        height=500,
        margin=dict(l=60, r=60, t=80, b=60),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        xaxis=dict(tickfont=dict(size=14)),
        yaxis=dict(tickfont=dict(size=14))
    )
    
    st.plotly_chart(fig_distribution, use_container_width=True, config={'displayModeBar': False})
    
    # Second chart - share distribution comparison (side by side pie charts)
    fig_pie = make_subplots(
        rows=1, cols=2, 
        specs=[[{'type':'domain'}, {'type':'domain'}]],
        subplot_titles=('Previous Period Share', 'Current Period Share')
    )
    
    # Previous period pie
    fig_pie.add_trace(go.Pie(
        labels=bucket_labels,
        values=previous_values,
        name="Previous",
        marker_colors=['#ff9999', '#66b3ff', '#99ff99', '#ffcc99'],
        textinfo='label+percent',
        textfont=dict(size=12)
    ), 1, 1)
    
    # Current period pie
    fig_pie.add_trace(go.Pie(
        labels=bucket_labels,
        values=current_values,
        name="Current",
        marker_colors=['#ff6666', '#3399ff', '#66ff66', '#ffb366'],
        textinfo='label+percent',
        textfont=dict(size=12)
    ), 1, 2)
    
    fig_pie.update_layout(
        title=dict(text="Ranking Distribution Share Comparison", font=dict(size=20)),
        height=500,
        margin=dict(l=60, r=60, t=100, b=60),
        paper_bgcolor='rgba(0,0,0,0)'
    )
    
    st.plotly_chart(fig_pie, use_container_width=True, config={'displayModeBar': False})
    
    # Detailed changes table
    st.markdown('<div class="section-header">📋 Detailed Changes by Ranking Bucket</div>', unsafe_allow_html=True)
    
    # Create detailed table
    table_data = []
    bucket_map = {
        'top_3': 'Top 3 (1-3)',
        'top_4_10': '4-10',
        'top_11_20': '11-20',
        'top_21_plus': '21+'
    }
    
    for bucket_key, bucket_name in bucket_map.items():
        data = results['bucket_changes'][bucket_key]
        table_data.append({
            'Ranking Position': bucket_name,
            'Previous Count': data['previous'],
            'Current Count': data['current'],
            'Change': data['change'],
            'Change %': f"{data['change_pct']:.1f}%",
            'Previous Share': f"{data['previous_share']:.1f}%",
            'Current Share': f"{data['current_share']:.1f}%"
        })
    
    table_df = pd.DataFrame(table_data)
    st.dataframe(table_df, use_container_width=True)
    
    # Strategic insights - AI-powered
    st.markdown('<div class="section-header">💡 Strategic Insights & Interpretation</div>', unsafe_allow_html=True)
    
    st.info("💡 **AI-Powered Insights:** Strategic insights are generated using AI. Set your OpenAI API key as an environment variable `OPENAI_API_KEY` or in Streamlit secrets.")
    
    with st.spinner("🤖 Generating AI-powered strategic insights..."):
        analysis_summary = create_visibility_analysis_summary(results)
        insights = generate_chatgpt_insights(analysis_summary, "visibility")
        st.markdown(f'<div class="insight-box">{insights.replace(chr(10), "<br>")}</div>', unsafe_allow_html=True)
    
    # Download section
    st.markdown('<div class="section-header">📥 Download Results</div>', unsafe_allow_html=True)
    
    # Create summary report
    summary_report = create_visibility_summary_report(results)
    
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            label="📄 Download Summary Report",
            data=summary_report,
            file_name=f"keyword_visibility_analysis_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain"
        )
    
    with col2:
        # Convert table to CSV for download
        csv_buffer = io.StringIO()
        table_df.to_csv(csv_buffer, index=False)
        st.download_button(
            label="📊 Download Data Table (CSV)",
            data=csv_buffer.getvalue(),
            file_name=f"keyword_visibility_data_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv"
        )

def generate_visibility_insights(results):
    """Generate strategic, actionable insights based on the visibility analysis"""
    
    total_change = results['total_change']
    total_change_pct = results['total_change_pct']
    
    top_3_change = results['bucket_changes']['top_3']['change_pct']
    top_3_current = results['bucket_changes']['top_3']['current_share']
    top_4_10_current = results['bucket_changes']['top_4_10']['current_share']
    top_10_current_share = top_3_current + top_4_10_current
    top_10_previous_share = (results['bucket_changes']['top_3']['previous_share'] + 
                            results['bucket_changes']['top_4_10']['previous_share'])
    top_21_plus_share = results['bucket_changes']['top_21_plus']['current_share']
    
    insights = []
    priority_actions = []
    
    # Overall trend analysis with actionable context
    if total_change > 100:
        insights.append(f"<b>🟢 Strong Keyword Expansion:</b> You've gained {abs(total_change):,} new keyword rankings ({total_change_pct:+.1f}%), indicating successful content expansion and improved domain authority.")
        priority_actions.append("Identify which content types/keywords drove this growth and double down on those strategies")
    elif total_change > 0:
        insights.append(f"<b>🟢 Moderate Growth:</b> You're ranking for {abs(total_change):,} more keywords ({total_change_pct:+.1f}%). This steady growth suggests sustainable SEO progress.")
    elif total_change < -100:
        insights.append(f"<b>🔴 Significant Keyword Loss:</b> You've lost rankings for {abs(total_change):,} keywords ({total_change_pct:.1f}%). This may indicate technical issues, content quality decline, or increased competition.")
        priority_actions.append("Audit lost keywords to identify patterns - check for technical SEO issues, content updates needed, or competitive displacement")
    elif total_change < 0:
        insights.append(f"<b>🟡 Minor Decline:</b> Lost {abs(total_change):,} keyword rankings ({total_change_pct:.1f}%). Monitor closely - could be seasonal or competitive shifts.")
    else:
        insights.append("<b>🟨 Stable Footprint:</b> Keyword count remained stable. Focus on improving positions rather than expanding breadth.")
    
    # Quality analysis with specific recommendations
    if top_3_change > 15:
        insights.append(f"<b>🟢 Exceptional Authority Growth:</b> Top 3 rankings increased by {top_3_change:.1f}% (now {top_3_current:.1f}% of total). This indicates strong content quality and E-E-A-T signals.")
        priority_actions.append("Analyze top 3 keywords to identify content patterns and replicate success across similar topics")
    elif top_3_change > 5:
        insights.append(f"<b>🟢 Strong Quality Improvement:</b> Top 3 rankings improved by {top_3_change:.1f}% (now {top_3_current:.1f}% of total). Your content relevance and authority are strengthening.")
    elif top_3_change > 0:
        insights.append(f"<b>🟢 Positive Momentum:</b> Top 3 rankings increased by {top_3_change:.1f}% (now {top_3_current:.1f}% of total). Continue focusing on high-intent, conversion-focused keywords.")
    elif top_3_change < -15:
        insights.append(f"<b>🔴 Critical Authority Decline:</b> Top 3 rankings dropped by {abs(top_3_change):.1f}% (now {top_3_current:.1f}% of total). This suggests serious competitive pressure or content quality issues.")
        priority_actions.append("URGENT: Review top 3 lost keywords - check for content freshness, backlink loss, or algorithm updates affecting these terms")
    elif top_3_change < -5:
        insights.append(f"<b>🔴 Authority Concern:</b> Top 3 rankings declined by {abs(top_3_change):.1f}% (now {top_3_current:.1f}% of total). Investigate content updates, competitor activity, and technical SEO.")
        priority_actions.append("Review pages that lost top 3 positions - update content, improve internal linking, and check for technical issues")
    
    # Strategic positioning analysis
    if total_change < 0 and top_10_current_share > top_10_previous_share + 5:
        insights.append(f"<b>🎯 Quality Over Quantity Success:</b> Despite losing {abs(total_change):,} total keywords, your top 10 share increased from {top_10_previous_share:.1f}% to {top_10_current_share:.1f}%. This indicates successful focus on high-value terms.")
        priority_actions.append("Continue this quality-focused strategy - prioritize ranking improvements over keyword expansion")
    elif total_change > 0 and top_10_current_share < top_10_previous_share - 5:
        insights.append(f"<b>⚠️ Breadth vs Depth Trade-off:</b> You gained {total_change:,} keywords but top 10 share dropped from {top_10_previous_share:.1f}% to {top_10_current_share:.1f}%. Consider consolidating efforts on improving positions for existing keywords.")
        priority_actions.append("Focus on improving positions for current rankings rather than expanding to new keywords - prioritize top 20 keywords for optimization")
    
    # Deep-dive opportunities
    if top_21_plus_share > 50:
        insights.append(f"<b>🎯 Major Optimization Opportunity:</b> {top_21_plus_share:.1f}% of keywords rank beyond position 20. These represent significant untapped potential - even moving to position 10 could drive substantial traffic.")
        priority_actions.append(f"Create optimization plan for top {int(results['total_current'] * 0.2):,} keywords ranking 21-50 - focus on on-page SEO, content enhancement, and internal linking")
    elif top_21_plus_share > 40:
        insights.append(f"<b>🎯 Optimization Opportunity:</b> {top_21_plus_share:.1f}% of keywords rank beyond position 20. Focus on improving on-page SEO, building topic authority, and enhancing content depth for these terms.")
        priority_actions.append("Identify keywords ranking 21-30 with high search volume and optimize those pages for quick wins")
    
    # Competitive positioning
    if top_10_current_share < 30:
        insights.append(f"<b>⚠️ Low Top 10 Concentration:</b> Only {top_10_current_share:.1f}% of keywords are in top 10. This suggests you're spread thin across many keywords. Consider focusing on fewer, higher-value terms.")
        priority_actions.append("Audit keyword portfolio - identify 50-100 highest-value keywords and prioritize optimization efforts on those")
    
    # Build final insights with priority actions
    if priority_actions:
        insights.append(f"<b>📋 Priority Actions:</b><ul>{''.join([f'<li>{action}</li>' for action in priority_actions[:3]])}</ul>")
    
    return "<br><br>".join(insights)

def create_visibility_summary_report(results):
    """Create downloadable visibility analysis report"""
    
    report = f"""
KEYWORD VISIBILITY ANALYSIS REPORT
Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

===========================================
EXECUTIVE SUMMARY
===========================================

Total Keywords (Current): {results['total_current']:,}
Total Keywords (Previous): {results['total_previous']:,}
Net Change: {results['total_change']:,} ({results['total_change_pct']:+.1f}%)

===========================================
RANKING DISTRIBUTION ANALYSIS
===========================================

Current Period Distribution:
"""
    
    for bucket, data in results['bucket_changes'].items():
        report += f"• {bucket.replace('_', ' ').title()}: {data['current']:,} keywords ({data['current_share']:.1f}%)\n"
    
    report += f"""

Previous Period Distribution:
"""
    
    for bucket, data in results['bucket_changes'].items():
        report += f"• {bucket.replace('_', ' ').title()}: {data['previous']:,} keywords ({data['previous_share']:.1f}%)\n"
    
    report += f"""

Changes by Bucket:
"""
    
    for bucket, data in results['bucket_changes'].items():
        change = data['change']
        change_pct = data['change_pct']
        report += f"• {bucket.replace('_', ' ').title()}: {change:+,} ({change_pct:+.1f}%)\n"
    
    report += f"""

===========================================
STRATEGIC INSIGHTS
===========================================

{generate_visibility_insights(results).replace('<b>', '').replace('</b>', '').replace('<br><br>', '\n\n').replace('🟢', '• ').replace('🟡', '• ').replace('🔴', '• ').replace('🎯', '• ').replace('⚠️', '• ').replace('🔧', '• ')}

===========================================
"""
    
    return report

def analyze_performance_pattern(clicks_delta, impr_delta, ctr_delta_pp, sessions_delta=None):
    """Analyze performance pattern from metrics changes"""
    
    pattern = {
        'description': '',
        'detail': '',
        'severity': 'neutral'
    }
    
    # Determine pattern based on metric changes
    if clicks_delta > 0 and impr_delta > 0 and ctr_delta_pp > 0:
        pattern['description'] = "Strong Growth Pattern"
        pattern['detail'] = "All key metrics are improving: clicks, impressions, and CTR are all up. This indicates successful SEO efforts with both visibility and engagement gains."
        pattern['severity'] = 'positive'
    elif clicks_delta > 0 and impr_delta > 0 and ctr_delta_pp < 0:
        pattern['description'] = "Visibility Growth, CTR Decline"
        pattern['detail'] = "Clicks and impressions are up, but CTR is down. This suggests increased visibility but potential issues with result relevance or SERP feature competition."
        pattern['severity'] = 'warning'
    elif clicks_delta < 0 and impr_delta < 0 and ctr_delta_pp < 0:
        pattern['description'] = "Declining Performance"
        pattern['detail'] = "All metrics are declining. This indicates a need for immediate attention to content quality, technical SEO, or competitive positioning."
        pattern['severity'] = 'negative'
    elif clicks_delta > 0 and impr_delta < 0:
        pattern['description'] = "Efficiency Improvement"
        pattern['detail'] = "Clicks are up despite fewer impressions, indicating better targeting and relevance. CTR improvement is driving results."
        pattern['severity'] = 'positive'
    elif clicks_delta < 0 and impr_delta > 0:
        pattern['description'] = "Visibility Up, Engagement Down"
        pattern['detail'] = "More impressions but fewer clicks suggests ranking for less relevant queries or increased SERP feature competition."
        pattern['severity'] = 'warning'
    else:
        pattern['description'] = "Mixed Performance"
        pattern['detail'] = "Metrics show mixed signals. Review individual query and page performance to identify specific opportunities."
        pattern['severity'] = 'neutral'
    
    return pattern

def data_export_instructions():
    """Comprehensive guide for exporting data from various SEO tools"""

    st.markdown("### 📋 Data Export Guide")
    st.caption(
        "Step-by-step instructions for exporting the exact files you’ll need from **Semrush**, **Google Search Console**, and **GA4**."
    )

    st.markdown(
        """
        <div style="
            background-color:#f9f9f9;
            border:1px solid #e6e6e6;
            border-radius:12px;
            padding:18px;
        ">
            <h4 style="margin-top:0;">🎯 Quick Reference</h4>
            <p style="margin-bottom:0;">
                • Export files in <b>CSV</b> or <b>Excel</b> format (never PDF).<br>
                • Use consistent naming: <code>client_tool_report_period.csv</code><br>
                • Keep date ranges consistent across all exports.<br>
                • For year-over-year, compare the same month from the previous year.
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.markdown("<div style='margin-top:30px;'></div>", unsafe_allow_html=True)
    
    # Create expandable sections for each tool
    with st.expander("🔍 **Semrush Exports** - Keyword Rankings & Competition Data", expanded=True):
        
        st.markdown("### Organic Research → Positions")
        st.markdown("""
        **Use for:** Keyword Visibility Trends, Competitor Analysis
        
        **Steps:**
        1. Go to **Domain Analytics → Organic Research → Positions**
        2. Set **Database** (e.g., United States)
        3. Set **Device = Desktop** 
        4. Set **Date = current month** (for current data)
        5. Click **Export → CSV** or **Excel**
        6. Save as: `client_semrush_positions_YYYY-MM.csv`
        
        **For Year-over-Year Analysis:**
        - Repeat above steps but change **Date** to same month last year
        - Save as: `client_semrush_positions_YYYY-MM_LY.csv`
        """)
        
        st.markdown("### Organic Research → Position Changes")
        st.markdown("""
        **Use for:** Keyword Movement Analysis
        
        **Steps:**
        1. Go to **Organic Research → Position Changes**
        2. Set **Range = Last 12 months** (or desired period)
        3. Click **Export → CSV** or **Excel**
        4. Save as: `client_semrush_position-changes_last12m.csv`
        """)
        
        st.markdown("### Organic Research → Pages")
        st.markdown("""
        **Use for:** Page Performance Analysis
        
        **Steps:**
        1. Go to **Organic Research → Pages**
        2. Set **Date = current month**
        3. Click **Export → CSV** or **Excel**
        4. Save as: `client_semrush_pages_YYYY-MM.csv`
        """)
        
        st.markdown("### Organic Research → Competitors")
        st.markdown("""
        **Use for:** Competitor Gap Analysis
        
        **Steps:**
        1. Go to **Organic Research → Competitors**
        2. Click **Export → CSV** or **Excel**
        3. Save as: `client_semrush_competitors_YYYY-MM.csv`
        """)
        
        st.markdown("### Position Tracking (Optional)")
        st.markdown("""
        **Use for:** SERP Features Analysis (if project exists)
        
        **Steps:**
        1. Go to **Projects → Position Tracking → Overview**
        2. If **AI Overviews** enabled, export that tab too
        3. Click **Export → CSV** or **Excel**
        4. Save as: `client_semrush_tracking_YYYY-MM.csv`
        
        ⚠️ **Note:** Only available if you have an existing Position Tracking project
        """)

    with st.expander("🔎 **Google Search Console Exports** - Click & Impression Data"):
        
        st.markdown("### Search Results - Sitewide Compare")
        st.markdown("""
        **Use for:** Traffic Attribution Analysis
        
        **Steps:**
        1. Go to **Search results** in left navigation
        2. Set top filters:
           - **Search type = Web**
           - **Date → Compare → Last 3 months vs Same period last year → Apply**
        3. Click **Export** (top right) → **CSV** or **Excel**
        4. Save as: `client_gsc_search-results_compare_[dates].csv`
        """)
        
        st.markdown("### Queries - Compare View")
        st.markdown("""
        **Use for:** Query Performance Analysis
        
        **Steps:**
        1. In **Search results**, click **Queries** tab
        2. Ensure **Compare** date setting is still applied from previous step
        3. Click **Export → CSV** or **Excel**
        4. Save as: `client_gsc_queries_compare_[dates].csv`
        """)
        
        st.markdown("### Pages - Compare View")
        st.markdown("""
        **Use for:** Page Performance Analysis (GSC data)
        
        **Steps:**
        1. In **Search results**, click **Pages** tab  
        2. Ensure **Compare** date setting is still applied
        3. Click **Export → CSV** or **Excel**
        4. Save as: `client_gsc_pages_compare_[dates].csv`
        """)
        
        st.markdown("""
        💡 **Pro Tip:** Set up your date comparison once in Search Results, then export all three views (sitewide, queries, pages) with the same date settings for consistency.
        """)

    with st.expander("📊 **GA4 Exports** - Traffic & Conversion Data"):
        
        st.markdown("### Traffic Acquisition - Organic Search Only")
        st.markdown("""
        **Use for:** Traffic Attribution Analysis
        
        **Steps:**
        1. Go to **Reports → Acquisition → Traffic acquisition**
        2. Set **date picker** to same window used in GSC
           - For YoY analysis, use **Compare** to previous year
        3. **Add filter:** Session default channel group = **Organic Search**
        4. Click **Share this report → Download file → CSV**
        5. Save as: `client_ga4_traffic-acquisition_organic_[dates].csv`
        """)
        
        st.markdown("### Landing Page - Organic Search Only")
        st.markdown("""
        **Use for:** Conversion Optimization Analysis
        
        **Steps:**
        1. Go to **Reports → Engagement → Landing page**
        2. Set **date picker** to same window as above
        3. **Add filter:** Session default channel group = **Organic Search**  
        4. Click **Share this report → Download file → CSV**
        5. Save as: `client_ga4_landing-page_organic_[dates].csv`
        """)
        
        st.markdown("### If Landing Page Report Missing")
        st.markdown("""
        **Alternative: Build Custom Exploration**
        
        1. Go to **Explore** → **+ Blank exploration**
        2. **Name it:** Landing Page -- Organic Search
        3. **Add dimension:** Landing page + query string
        4. **Add metrics:** Sessions, Active users, Engaged sessions, Average engagement time, Key events
        5. **Add filter:** Session default channel group = "Organic Search"
        6. **Export:** Top right → Export → CSV
        
        This creates the same data as the standard Landing Page report.
        """)

    with st.expander("✅ **File Validation Checklist**"):
        st.markdown("""
        Before uploading your files, verify:
        
        **✅ Format Requirements:**
        - Files are in CSV or Excel format (never PDF)
        - Column headers are present and readable
        - Data contains expected metrics (clicks, impressions, positions, etc.)
        
        **✅ Date Consistency:**
        - All files use the same date ranges
        - Year-over-year comparisons use same month from previous year
        - Compare periods match between GSC and GA4 exports
        
        **✅ File Naming:**
        - Clear, consistent naming convention
        - Include client, tool, report type, and date period
        - Example: `acme_semrush_positions_2024-08.csv`
        
        **✅ Data Quality:**
        - Files contain data (not empty exports)
        - Position data is numeric (not text)
        - Keyword/Query lists are complete
        """)

    with st.expander("🆘 **Troubleshooting Common Issues**"):
        st.markdown("""
        **"Column not found" errors:**
        - Column names vary between exports - the tool auto-detects most variations
        - Ensure you're exporting from the correct section (Positions vs Position Changes)
        - Check that headers are in English (not localized)
        
        **Excel files won't upload:**
        - Try exporting as CSV instead
        - Ensure file isn't password protected
        - Check file size (very large files may timeout)
        
        **Missing data in analysis:**
        - Verify date ranges match between current and previous exports
        - Some tools have data limitations (GSC keeps ~16 months)
        - Position Tracking requires existing project setup
        
        **GSC Compare issues:**
        - Use "Same period last year" not custom date ranges
        - Ensure you're comparing like periods (3 months vs 3 months)
        - Web search type should be selected
        
        **GA4 Filter problems:**
        - Use exact text "Organic Search" for channel group filter
        - Default channel group is different from source/medium
        - Some accounts have custom channel definitions
        """)

    st.markdown("---")
    st.success("💡 **Ready to start?** Choose an analysis tab above and follow the specific file requirements for each section!")

def keyword_movement_analysis():
    """Analyze keyword movement distribution from Semrush Position Changes"""
    st.markdown('<div class="section-header">🔄 Keyword Movement Distribution</div>', unsafe_allow_html=True)
    
    # Modern instruction design using containers and columns
    with st.container():
        st.markdown("### 📈 Analysis Overview")
        
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown("""
            This analysis examines how your keyword rankings changed over time to identify:
            
            **🎯 Key Questions Answered:**
            - Which keywords are improving vs declining?
            - What's your overall momentum (improved:declined ratio)?
            - Where are keywords flowing between ranking buckets?
            - Which specific keywords need attention?
            """)
        
        with col2:
            st.info("""
            **💡 Strategic Value**
            
            Identifies momentum trends and specific optimization opportunities for your keyword portfolio.
            """)
    
    # File requirements in expandable section
    with st.expander("📁 **File Requirements & Setup**", expanded=False):
        st.markdown("""
        **Required Files:** 1 Semrush Position Changes export
        
        | Setting | Requirement |
        |---------|-------------|
        | **Export From** | Organic Research → Position Changes |
        | **Time Period** | Last 12 months (recommended) |
        | **Format** | CSV or Excel |
        | **Must Include** | Keyword, Position, Previous Position |
        
        **🔍 Methodology Note:**
        - Position = 0 treated as "not ranked" (worst position)
        - Falling out (→0) = Declined, Newly ranked (0→#) = Improved
        - Movement = Previous Position - Position (positive = improvement)
        """)
    
    # Key insights preview
    st.markdown("### 🎯 Analysis Insights You'll Get")
    
    insight_col1, insight_col2, insight_col3, insight_col4 = st.columns(4)
    
    with insight_col1:
        st.markdown("""
        **📊 Movement Distribution**
        - Improved/Declined/Unchanged counts
        - Overall momentum ratio
        """)
    
    with insight_col2:
        st.markdown("""
        **🏆 Top Winners**
        - Biggest ranking improvements
        - New #1 rankings priority
        """)
    
    with insight_col3:
        st.markdown("""
        **📉 Top Losers**
        - Keywords needing attention
        - Decline root cause analysis
        """)
    
    with insight_col4:
        st.markdown("""
        **🔄 Bucket Flow**
        - Ranking tier movements
        - Sources of elite rankings
        """)
    
    st.markdown("---")
    
    # File upload section
    st.markdown("### 📤 Upload Your Data File")
    
    position_changes_file = st.file_uploader(
        "Upload Semrush Position Changes file",
        type=['csv', 'xlsx', 'xls'],
        key="position_changes",
        help="Export from Semrush: Organic Research → Position Changes (CSV or Excel format)"
    )
    
    # Process file if uploaded
    if position_changes_file is not None:
        # Add Run Analysis button (centered)
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            run_movement_analysis = st.button("🚀 Run Movement Analysis", key="run_movement", type="primary", use_container_width=True)
        
        # Display results outside column context for full width
        if run_movement_analysis:
            with st.spinner("🔄 Analyzing keyword movements..."):
                try:
                    # Load and validate data
                    df = normalize_columns(read_uploaded_file(position_changes_file))
                    
                    # Validate required columns
                    validation_passed, validation_message = validate_movement_data(df)
                    
                    if not validation_passed:
                        st.error(validation_message)
                        st.stop()
                    
                    # Perform analysis
                    movement_results = analyze_keyword_movement(df)
                    
                    # Display results - NOW IN FULL WIDTH
                    display_movement_results(movement_results)
                    
                except Exception as e:
                    st.error(f"❌ Error processing file: {str(e)}")
                    st.info("💡 Please ensure you've uploaded a valid Semrush Position Changes file")
    else:
        st.info("📤 Please upload a Semrush Position Changes file to begin analysis")

def validate_movement_data(df):
    """Validate the Position Changes data"""
    required_columns = ['Keyword', 'Position', 'Previous position']
    
    # Find columns using flexible matching
    kw_col = find_column(df.columns, ['keyword'])
    pos_col = find_column(df.columns, ['position']) 
    prev_pos_col = find_column(df.columns, ['previous position', 'prev position', 'previous'])
    
    missing_columns = []
    if not kw_col:
        missing_columns.append('Keyword')
    if not pos_col:
        missing_columns.append('Position')  
    if not prev_pos_col:
        missing_columns.append('Previous Position')
    
    if missing_columns:
        return False, f"❌ Missing required columns: {missing_columns}. Available columns: {list(df.columns)[:10]}"
    
    # Check if data is not empty
    if len(df) == 0:
        return False, "❌ File appears to be empty"
    
    return True, "✅ Data validation passed"

def analyze_keyword_movement(df):
    """Analyze keyword movement patterns"""
    
    # Find and rename columns
    kw_col = find_column(df.columns, ['keyword'])
    pos_col = find_column(df.columns, ['position'])
    prev_pos_col = find_column(df.columns, ['previous position', 'prev position', 'previous'])
    url_col = find_column(df.columns, ['url', 'page', 'landing page'])
    
    # Prepare working dataframe
    work_df = pd.DataFrame()
    work_df['Keyword'] = df[kw_col].astype(str).str.strip()
    work_df['Position'] = pd.to_numeric(df[pos_col], errors='coerce')
    work_df['Previous_Position'] = pd.to_numeric(df[prev_pos_col], errors='coerce')
    
    if url_col:
        work_df['URL'] = df[url_col].astype(str)
    
    # Remove rows with missing position data
    work_df = work_df.dropna(subset=['Position', 'Previous_Position'])
    
    # Calculate movement (real numbers for tables)
    work_df['Movement'] = work_df['Previous_Position'] - work_df['Position']  # positive = improved rank
    
    # Status classification treating 0 as "not ranked" (worst)
    def effective_rank(series):
        return np.where((series <= 0) | pd.isna(series), 1000, series)
    
    eff_prev = effective_rank(work_df['Previous_Position'])
    eff_now = effective_rank(work_df['Position'])
    status_movement = eff_prev - eff_now  # >0 improved, <0 declined, =0 unchanged
    
    work_df['Status'] = np.where(status_movement > 0, 'Improved',
                        np.where(status_movement < 0, 'Declined', 'Unchanged'))
    
    # Movement distribution
    counts = work_df['Status'].value_counts().reindex(['Improved', 'Declined', 'Unchanged']).fillna(0).astype(int)
    improved, declined, unchanged = int(counts.get('Improved', 0)), int(counts.get('Declined', 0)), int(counts.get('Unchanged', 0))
    ratio = (improved / declined) if declined > 0 else np.inf
    
    # Top improvers (exclude Position=0, prioritize #1)
    improved_df = work_df[(work_df['Movement'] > 0) & (work_df['Position'] > 0)].copy()
    improved_df['Reached_1'] = (improved_df['Position'] == 1).astype(int)
    top_improvers = improved_df.sort_values(['Reached_1', 'Movement'], ascending=[False, False]).head(25)
    
    # Top decliners (exclude Previous_Position=0, keep drops to 0)
    declined_df = work_df[(work_df['Movement'] < 0) & (work_df['Previous_Position'] > 0)].copy()
    top_decliners = declined_df.sort_values('Movement', ascending=True).head(25)
    
    # Bucket analysis
    def bucketize_position(series):
        return pd.cut(series, bins=[-np.inf, 0, 3, 10, 20, np.inf], 
                     labels=['Invalid', 'Top 1-3', '4-10', '11-20', '21+'], right=True)
    
    bucket_order = ['Top 1-3', '4-10', '11-20', '21+']
    work_df['Prev_Bucket'] = bucketize_position(work_df['Previous_Position'])
    work_df['Now_Bucket'] = bucketize_position(work_df['Position'])
    
    # Filter valid buckets for transition analysis
    bucket_df = work_df[(work_df['Prev_Bucket'] != 'Invalid') & (work_df['Now_Bucket'] != 'Invalid')].copy()
    
    # Transition matrix
    transition_matrix = pd.crosstab(bucket_df['Prev_Bucket'], bucket_df['Now_Bucket']).reindex(index=bucket_order, columns=bucket_order, fill_value=0)
    
    # Net flow analysis
    diagonal = pd.Series({b: transition_matrix.at[b, b] if (b in transition_matrix.index and b in transition_matrix.columns) else 0 for b in bucket_order})
    inflow = transition_matrix.sum(axis=0) - diagonal
    outflow = transition_matrix.sum(axis=1) - diagonal
    net_flow = inflow.reindex(bucket_order) - outflow.reindex(bucket_order)
    
    prev_counts = transition_matrix.sum(axis=1).reindex(bucket_order)
    now_counts = transition_matrix.sum(axis=0).reindex(bucket_order)
    delta_counts = now_counts - prev_counts
    
    # Sources of new Top 1-3
    new_top3 = bucket_df[bucket_df['Now_Bucket'] == 'Top 1-3']
    top3_sources = new_top3['Prev_Bucket'].value_counts().reindex(bucket_order, fill_value=0)
    
    return {
        'total_keywords': len(work_df),
        'movement_counts': {'improved': improved, 'declined': declined, 'unchanged': unchanged},
        'ratio': ratio,
        'top_improvers': top_improvers,
        'top_decliners': top_decliners,
        'transition_matrix': transition_matrix,
        'bucket_flow': {
            'prev_counts': prev_counts,
            'now_counts': now_counts, 
            'delta_counts': delta_counts,
            'inflow': inflow.reindex(bucket_order),
            'outflow': outflow.reindex(bucket_order),
            'net_flow': net_flow
        },
        'top3_sources': top3_sources,
        'raw_data': work_df
    }

def display_movement_results(results):
    """Display keyword movement analysis results"""
    
    # Key metrics
    st.markdown('<div class="section-header">📈 Movement Distribution Summary</div>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            label="Keywords Analyzed",
            value=f"{results['total_keywords']:,}"
        )
    
    with col2:
        st.metric(
            label="Improved Rankings",
            value=f"{results['movement_counts']['improved']:,}",
            delta=f"+{results['movement_counts']['improved']:,} keywords"
        )
    
    with col3:
        # For declined rankings: negative is bad, should show red
        # Using "normal" so negative delta shows as red (bad)
        st.metric(
            label="Declined Rankings", 
            value=f"{results['movement_counts']['declined']:,}",
            delta=f"-{results['movement_counts']['declined']:,} keywords",
            delta_color="normal"  # Normal mode: negative = red (bad), positive = green (good)
        )
    
    with col4:
        ratio_display = f"{results['ratio']:.2f}" if results['ratio'] != np.inf else "∞"
        st.metric(
            label="Improved:Declined Ratio",
            value=ratio_display,
            help="Higher ratio indicates more keywords improving than declining"
        )
    
    # Distribution chart
    st.markdown('<div class="section-header">📊 Movement Distribution</div>', unsafe_allow_html=True)
    
    # Full-width bar chart
    dist_data = results['movement_counts']
    fig_dist = go.Figure(data=[
        go.Bar(x=list(dist_data.keys()), 
               y=list(dist_data.values()),
               marker_color=['#2ecc71', '#e74c3c', '#95a5a6'],
               text=[f"{val:,}" for val in dist_data.values()],
               textposition='auto',
               textfont=dict(size=16, color='white')
        )
    ])
    
    fig_dist.update_layout(
        title=dict(text='Keyword Movement Distribution', font=dict(size=20)),
        xaxis_title='Movement Type',
        yaxis_title='Number of Keywords',
        height=500,
        margin=dict(l=60, r=60, t=80, b=60),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        xaxis=dict(tickfont=dict(size=14)),
        yaxis=dict(tickfont=dict(size=14))
    )
    
    st.plotly_chart(fig_dist, use_container_width=True, config={'displayModeBar': False})
    
    # Full-width pie chart 
    labels = list(results['movement_counts'].keys())
    values = list(results['movement_counts'].values())
    
    fig_pie = go.Figure(data=[go.Pie(
        labels=labels,
        values=values,
        marker_colors=['#2ecc71', '#e74c3c', '#95a5a6'],
        textinfo='label+percent+value',
        textfont=dict(size=14),
        pull=[0.05, 0.05, 0],
        hole=0.3
    )])
    
    fig_pie.update_layout(
        title=dict(text='Movement Distribution Share', font=dict(size=20)),
        height=500,
        margin=dict(l=60, r=60, t=80, b=60),
        paper_bgcolor='rgba(0,0,0,0)',
        annotations=[dict(text=f'Total<br>{sum(values):,}', x=0.5, y=0.5, font_size=16, showarrow=False)]
    )
    
    st.plotly_chart(fig_pie, use_container_width=True, config={'displayModeBar': False})
    
    # Top winners and losers - full width layout
    st.markdown('<div class="section-header">🏆 Top Moving Keywords</div>', unsafe_allow_html=True)
    
    # Use tabs for winners/losers to save vertical space and improve readability
    winner_tab, loser_tab = st.tabs(["📈 Top Improving Keywords", "📉 Top Declining Keywords"])
    
    with winner_tab:
        st.markdown("*Keywords with biggest positive movement (prioritizing new #1 rankings)*")
        
        if not results['top_improvers'].empty:
            display_cols = ['Keyword', 'Previous_Position', 'Position', 'Movement']
            if 'URL' in results['top_improvers'].columns:
                display_cols.append('URL')
            
            improvers_display = results['top_improvers'][display_cols].head(20).copy()
            improvers_display.columns = ['Keyword', 'Previous Pos', 'Current Pos', 'Movement'] + (['URL'] if 'URL' in display_cols else [])
            st.dataframe(improvers_display, use_container_width=True, hide_index=True, height=400)
        else:
            st.info("No improving keywords found with the current criteria")
    
    with loser_tab:
        st.markdown("*Keywords with biggest negative movement (excluding newly ranked keywords)*")
        
        if not results['top_decliners'].empty:
            display_cols = ['Keyword', 'Previous_Position', 'Position', 'Movement']
            if 'URL' in results['top_decliners'].columns:
                display_cols.append('URL')
            
            decliners_display = results['top_decliners'][display_cols].head(20).copy()
            decliners_display.columns = ['Keyword', 'Previous Pos', 'Current Pos', 'Movement'] + (['URL'] if 'URL' in display_cols else [])
            st.dataframe(decliners_display, use_container_width=True, hide_index=True, height=400)
        else:
            st.info("No declining keywords found with the current criteria")
    
    # Bucket flow analysis
    st.markdown('<div class="section-header">🔄 Ranking Bucket Flow Analysis</div>', unsafe_allow_html=True)
    
    # Transition matrix
    st.markdown("**Transition Matrix - Previous Bucket → Current Bucket**")
    st.dataframe(results['transition_matrix'], use_container_width=True)
    
    # Net flow table
    st.markdown("**Net Movement by Ranking Bucket**")
    flow_table = pd.DataFrame({
        'Ranking Bucket': results['bucket_flow']['prev_counts'].index,
        'Previous Count': results['bucket_flow']['prev_counts'].values,
        'Current Count': results['bucket_flow']['now_counts'].values,
        'Change': results['bucket_flow']['delta_counts'].values,
        'Inflow': results['bucket_flow']['inflow'].values,
        'Outflow': results['bucket_flow']['outflow'].values,
        'Net Flow': results['bucket_flow']['net_flow'].values
    })
    
    st.dataframe(flow_table, use_container_width=True, hide_index=True)
    
    # Sources of new top 3 rankings
    if not results['top3_sources'].empty:
        st.markdown("**Sources of New Top 1-3 Rankings**")
        sources_df = pd.DataFrame({
            'Previous Bucket': results['top3_sources'].index,
            'Keywords Moved to Top 1-3': results['top3_sources'].values
        })
        st.dataframe(sources_df, use_container_width=True, hide_index=True)
    
    # Strategic insights - AI-powered
    st.markdown('<div class="section-header">💡 Strategic Insights</div>', unsafe_allow_html=True)
    
    st.info("💡 **AI-Powered Insights:** Strategic insights are generated using AI. Set your OpenAI API key as an environment variable `OPENAI_API_KEY` or in Streamlit secrets.")
    
    with st.spinner("🤖 Generating AI-powered strategic insights..."):
        analysis_summary = create_movement_analysis_summary(results)
        insights = generate_chatgpt_insights(analysis_summary, "movement")
        st.markdown(f'<div class="insight-box">{insights.replace(chr(10), "<br>")}</div>', unsafe_allow_html=True)
    
    # Download section
    st.markdown('<div class="section-header">📥 Download Results</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        summary_report = create_movement_summary_report(results)
        st.download_button(
            label="📄 Download Movement Analysis Report",
            data=summary_report,
            file_name=f"keyword_movement_analysis_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain"
        )
    
    with col2:
        # Convert top movers to CSV
        csv_buffer = io.StringIO()
        combined_movers = pd.concat([
            results['top_improvers'].head(15).assign(Type='Improver'),
            results['top_decliners'].head(15).assign(Type='Decliner')
        ])
        combined_movers.to_csv(csv_buffer, index=False)
        
        st.download_button(
            label="📊 Download Top Movers (CSV)",
            data=csv_buffer.getvalue(),
            file_name=f"keyword_top_movers_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv"
        )

def generate_movement_insights(results):
    """Generate strategic insights from movement analysis"""
    insights = []
    
    ratio = results['ratio']
    improved = results['movement_counts']['improved']
    declined = results['movement_counts']['declined']
    
    # Overall trend analysis
    if ratio > 1.5:
        insights.append(f"<b>🟢 Strong Upward Momentum:</b> With {improved:,} improving vs {declined:,} declining keywords (ratio: {ratio:.2f}), your overall keyword portfolio is strengthening significantly.")
    elif ratio > 1.0:
        insights.append(f"<b>🟢 Positive Trend:</b> More keywords improving ({improved:,}) than declining ({declined:,}), indicating healthy SEO momentum.")
    elif ratio > 0.8:
        insights.append(f"<b>🟡 Mixed Results:</b> Nearly balanced movement with {improved:,} improving vs {declined:,} declining. Focus on protecting top performers.")
    else:
        insights.append(f"<b>🔴 Declining Trend:</b> More keywords declining ({declined:,}) than improving ({improved:,}). Priority should be stabilizing rankings and identifying causes.")
    
    # Top 3 analysis
    if not results['top3_sources'].empty:
        top_source = results['top3_sources'].idxmax()
        top_count = results['top3_sources'].max()
        insights.append(f"<b>🎯 Top 3 Growth:</b> Most new top 3 rankings ({top_count}) came from the {top_source} bucket, showing your ability to push mid-performing keywords to elite positions.")
    
    # Bucket flow insights
    flow = results['bucket_flow']
    top3_net = flow['net_flow']['Top 1-3'] if 'Top 1-3' in flow['net_flow'].index else 0
    tail_net = flow['net_flow']['21+'] if '21+' in flow['net_flow'].index else 0
    
    if top3_net > 0 and tail_net < 0:
        insights.append("<b>🎯 Quality Consolidation:</b> You're successfully moving keywords from the long tail into top positions - a sign of content optimization working.")
    elif top3_net < 0:
        insights.append("<b>⚠️ Top Position Pressure:</b> You're losing some top 3 rankings. Review competitors and refresh your highest-value content.")
    
    # Actionable recommendations
    if results['movement_counts']['declined'] > 0:
        insights.append(f"<b>🔧 Action Items:</b> Review the {min(15, results['movement_counts']['declined'])} top declining keywords shown above. Look for content freshness, competitor analysis, and technical issues.")
    
    return "<br><br>".join(insights)

def create_movement_summary_report(results):
    """Create downloadable movement analysis report"""
    
    ratio_display = f"{results['ratio']:.2f}" if results['ratio'] != np.inf else "∞"
    
    report = f"""
KEYWORD MOVEMENT ANALYSIS REPORT
Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

===========================================
EXECUTIVE SUMMARY
===========================================

Total Keywords Analyzed: {results['total_keywords']:,}

Movement Distribution:
• Improved Rankings: {results['movement_counts']['improved']:,} keywords
• Declined Rankings: {results['movement_counts']['declined']:,} keywords  
• Unchanged Rankings: {results['movement_counts']['unchanged']:,} keywords
• Improved:Declined Ratio: {ratio_display}

===========================================
TOP IMPROVING KEYWORDS (Sample)
===========================================

"""
    
    if not results['top_improvers'].empty:
        for _, row in results['top_improvers'].head(10).iterrows():
            report += f"• {row['Keyword']} | {row['Previous_Position']} → {row['Position']} (+{row['Movement']:.0f})\n"
    
    report += f"""

===========================================
TOP DECLINING KEYWORDS (Sample)
===========================================

"""
    
    if not results['top_decliners'].empty:
        for _, row in results['top_decliners'].head(10).iterrows():
            report += f"• {row['Keyword']} | {row['Previous_Position']} → {row['Position']} ({row['Movement']:.0f})\n"
    
    report += f"""

===========================================
RANKING BUCKET FLOW ANALYSIS
===========================================

Bucket Changes:
"""
    
    for bucket in results['bucket_flow']['prev_counts'].index:
        prev_count = results['bucket_flow']['prev_counts'][bucket]
        now_count = results['bucket_flow']['now_counts'][bucket] 
        change = results['bucket_flow']['delta_counts'][bucket]
        report += f"• {bucket}: {prev_count} → {now_count} ({change:+})\n"
    
    report += f"""

===========================================
STRATEGIC INSIGHTS
===========================================

{generate_movement_insights(results).replace('<b>', '').replace('</b>', '').replace('<br><br>', '\n\n').replace('🟢', '• ').replace('🟡', '• ').replace('🔴', '• ').replace('🎯', '• ').replace('⚠️', '• ').replace('🔧', '• ')}

===========================================
"""
    
    return report

def page_performance_analysis():
    """Analyze page performance from Semrush and GSC data"""
    st.markdown('<div class="section-header">📄 Page Performance Analysis</div>', unsafe_allow_html=True)
    
    # Create sub-tabs for different data sources
    semrush_tab, gsc_tab = st.tabs(["🔍 Semrush Analysis", "📊 Google Search Console Analysis"])
    
    with semrush_tab:
        semrush_page_analysis()
    
    with gsc_tab:
        gsc_page_analysis()

def semrush_page_analysis():
    """Semrush Pages Analysis (original functionality)"""
    
    # Modern instruction design using containers and columns
    with st.container():
        st.markdown("### 📊 Semrush Analysis Overview")
        
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown("""
            This analysis examines your top-performing pages using Semrush data:
            
            **🎯 Key Questions Answered:**
            - How concentrated is your organic traffic? (Pareto analysis)
            - Which pages are most efficient at driving traffic per keyword?
            - Where do your content hubs generate the most value?
            - Which pages have untapped optimization potential?
            """)
        
        with col2:
            st.info("""
            **💡 Strategic Value**
            
            Identifies your highest-value pages for protection and reveals optimization opportunities.
            """)
    
    # File requirements in expandable section
    with st.expander("📁 **Semrush File Requirements**", expanded=False):
        st.markdown("""
        **Required Files:** 1 Semrush Pages export
        
        | Setting | Requirement |
        |---------|-------------|
        | **Export From** | Organic Research → Pages |
        | **Time Period** | Current month |
        | **Format** | CSV or Excel |
        | **Must Include** | URL, Traffic, Traffic %, Number of Keywords |
        """)
    
    st.markdown("---")
    
    # File upload section
    st.markdown("### 📤 Upload Semrush Pages Data")
    
    pages_file = st.file_uploader(
        "Upload Semrush Pages file",
        type=['csv', 'xlsx', 'xls'],
        key="semrush_pages",
        help="Export from Semrush: Organic Research → Pages"
    )
    
    # Process file if uploaded
    if pages_file is not None:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            run_semrush_analysis = st.button("🚀 Run Semrush Analysis", key="run_semrush_pages", type="primary", use_container_width=True)
        
        if run_semrush_analysis:
            with st.spinner("🔄 Analyzing Semrush page performance..."):
                try:
                    df = normalize_columns(read_uploaded_file(pages_file))
                    validation_passed, validation_message = validate_semrush_pages_data(df)
                    
                    if not validation_passed:
                        st.error(validation_message)
                        st.stop()
                    
                    pages_results = analyze_page_performance(df)
                    display_pages_results(pages_results)
                    
                except Exception as e:
                    st.error(f"❌ Error processing file: {str(e)}")
                    st.info("💡 Please ensure you've uploaded a valid Semrush Pages file")
    else:
        st.info("📤 Please upload a Semrush Pages file to begin analysis")

def validate_semrush_pages_data(df):
    """Validate the Semrush Pages data"""
    
    # Find columns using flexible matching
    url_col = find_column(df.columns, ['url', 'page', 'landing page'])
    traffic_col = find_column(df.columns, ['traffic'])
    
    missing_columns = []
    if not url_col:
        missing_columns.append('URL/Page')
    if not traffic_col:
        missing_columns.append('Traffic')
    
    if missing_columns:
        return False, f"❌ Missing required columns: {missing_columns}. Available columns: {list(df.columns)[:10]}"
    
    # Check if data is not empty
    if len(df) == 0:
        return False, "❌ File appears to be empty"
    
    return True, "✅ Data validation passed"

def gsc_page_analysis():
    """Google Search Console Pages Analysis"""
    
    with st.container():
        st.markdown("### 📊 Google Search Console Analysis Overview")
        
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown("""
            This analysis examines your page performance using GSC data:
            
            **🎯 Key Questions Answered:**
            - Which pages drive the most clicks and impressions?
            - What are your best and worst performing CTRs by page?
            - How do your pages perform across different countries?
            - Which pages have the best average positions?
            """)
        
        with col2:
            st.info("""
            **💡 Strategic Value**
            
            Shows actual search performance data and geographic distribution of your traffic.
            """)
    
    # File requirements in expandable section
    with st.expander("📁 **GSC File Requirements**", expanded=False):
        st.markdown("""
        **Required Files:** 1 GSC Pages Compare Excel file with multiple sheets
        
        | Sheet | Contains | Purpose |
        |-------|----------|---------|
        | **Pages** | Page performance data | Clicks, impressions, CTR, position by page |
        | **Countries** | Geographic data | Traffic distribution by country |
        
        **📋 Export Steps:**
        1. Go to Search Results in Google Search Console
        2. Set: Date → Compare → Last 3 months vs Same period last year
        3. Click Pages tab and export
        4. Add Countries data as additional sheet
        5. Save as Excel file with multiple sheets
        """)
    
    st.markdown("---")
    
    # File upload section
    st.markdown("### 📤 Upload GSC Pages Data")
    
    gsc_pages_file = st.file_uploader(
        "Upload GSC Pages Compare Excel file",
        type=['xlsx', 'xls'],
        key="gsc_pages_file",
        help="Excel file with Pages and Countries sheets"
    )
    
    # Process file if uploaded
    if gsc_pages_file is not None:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            run_gsc_analysis = st.button("🚀 Run GSC Analysis", key="run_gsc_pages", type="primary", use_container_width=True)
        
        if run_gsc_analysis:
            with st.spinner("🔄 Analyzing GSC page performance..."):
                try:
                    # Read Excel file with multiple sheets
                    pages_df, countries_df = read_gsc_excel_file(gsc_pages_file)
                    
                    # Validate data
                    validation_passed, validation_message = validate_gsc_pages_data(pages_df)
                    
                    if not validation_passed:
                        st.error(validation_message)
                        st.stop()
                    
                    # Perform analysis
                    gsc_results = analyze_gsc_page_performance(pages_df, countries_df)
                    
                    # Display results
                    display_gsc_results(gsc_results)
                    
                except Exception as e:
                    st.error(f"❌ Error processing file: {str(e)}")
                    st.info("💡 Please ensure you've uploaded a valid GSC Excel file with Pages and Countries sheets")
    else:
        st.info("📤 Please upload a GSC Pages Excel file to begin analysis")

def read_gsc_excel_file(uploaded_file):
    """Read GSC Excel file with Pages and Countries sheets"""
    
    try:
        # Read the Pages sheet (sheet 2, index 1)
        pages_df = pd.read_excel(uploaded_file, sheet_name=1)  # Sheet 2
        pages_df = normalize_columns(pages_df)
    except:
        # Fallback - try by name
        try:
            pages_df = pd.read_excel(uploaded_file, sheet_name='Pages')
            pages_df = normalize_columns(pages_df)
        except:
            raise ValueError("Could not find Pages sheet in Excel file")
    
    try:
        # Read the Countries sheet (sheet 3, index 2)
        countries_df = pd.read_excel(uploaded_file, sheet_name=2)  # Sheet 3
        countries_df = normalize_columns(countries_df)
    except:
        # Fallback - try by name
        try:
            countries_df = pd.read_excel(uploaded_file, sheet_name='Countries')
            countries_df = normalize_columns(countries_df)
        except:
            st.warning("Could not find Countries sheet - map will be skipped")
            countries_df = pd.DataFrame()
    
    return pages_df, countries_df

def validate_gsc_pages_data(df):
    """Validate GSC Pages data"""
    
    # Look for typical GSC columns
    page_col = find_column(df.columns, ['top pages', 'page', 'url'])
    clicks_col = find_column(df.columns, ['clicks'])
    
    missing_columns = []
    if not page_col:
        missing_columns.append('Page/URL')
    if not clicks_col:
        missing_columns.append('Clicks')
    
    if missing_columns:
        return False, f"❌ Missing required columns: {missing_columns}. Available columns: {list(df.columns)[:10]}"
    
    if len(df) == 0:
        return False, "❌ File appears to be empty"
    
    return True, "✅ Data validation passed"

def analyze_gsc_page_performance(pages_df, countries_df):
    """Analyze GSC page performance data"""
    
    # Find columns
    page_col = find_column(pages_df.columns, ['top pages', 'page', 'url'])
    clicks_now_col = find_column(pages_df.columns, ['last 3 months clicks', 'clicks'])
    clicks_prev_col = find_column(pages_df.columns, ['previous 3 months clicks', 'same period last year clicks'])
    impr_now_col = find_column(pages_df.columns, ['last 3 months impressions', 'impressions'])
    impr_prev_col = find_column(pages_df.columns, ['previous 3 months impressions', 'same period last year impressions'])
    ctr_now_col = find_column(pages_df.columns, ['last 3 months ctr', 'ctr'])
    position_col = find_column(pages_df.columns, ['position'])
    
    # Build working dataframe
    work_df = pd.DataFrame()
    work_df['Page'] = pages_df[page_col].astype(str).str.strip()
    
    # Normalize URLs to handle trailing slashes and case variations
    def normalize_url(url):
        if pd.isna(url) or url == '':
            return url
        url = str(url).strip()
        # Remove trailing slash (but keep if it's just the domain)
        if url.endswith('/') and url.count('/') > 2:
            url = url.rstrip('/')
        # Convert to lowercase for consistent comparison
        return url.lower()
    
    work_df['Page_Normalized'] = work_df['Page'].apply(normalize_url)
    work_df['Clicks_Now'] = pd.to_numeric(pages_df[clicks_now_col], errors='coerce') if clicks_now_col else 0
    work_df['Clicks_Prev'] = pd.to_numeric(pages_df[clicks_prev_col], errors='coerce') if clicks_prev_col else 0
    work_df['Impr_Now'] = pd.to_numeric(pages_df[impr_now_col], errors='coerce') if impr_now_col else 0
    work_df['Impr_Prev'] = pd.to_numeric(pages_df[impr_prev_col], errors='coerce') if impr_prev_col else 0
    
    if ctr_now_col:
        ctr_series = pages_df[ctr_now_col]
        if ctr_series.astype(str).str.contains('%').any():
            work_df['CTR_Now'] = pd.to_numeric(ctr_series.astype(str).str.replace('%', ''), errors='coerce')
        else:
            ctr_numeric = pd.to_numeric(ctr_series, errors='coerce')
            work_df['CTR_Now'] = ctr_numeric * 100 if ctr_numeric.max() <= 1.0 else ctr_numeric
    else:
        work_df['CTR_Now'] = np.where(work_df['Impr_Now'] > 0, work_df['Clicks_Now'] / work_df['Impr_Now'] * 100, 0)
    
    if position_col:
        work_df['Position'] = pd.to_numeric(pages_df[position_col], errors='coerce')
    else:
        work_df['Position'] = np.nan
    
    # Clean data
    work_df = work_df[work_df['Page'].notna() & work_df['Page'].ne('')].copy()
    
    # Aggregate duplicates by normalized URL (sum metrics, keep first URL for display)
    agg_df = work_df.groupby('Page_Normalized').agg({
        'Page': 'first',  # Keep first page URL for display
        'Clicks_Now': 'sum',  # Sum clicks for duplicates
        'Clicks_Prev': 'sum',  # Sum previous clicks
        'Impr_Now': 'sum',  # Sum impressions
        'Impr_Prev': 'sum',  # Sum previous impressions
        'Position': 'mean'  # Average position for duplicates
    }).reset_index()
    
    # Recalculate CTR after aggregation
    agg_df['CTR_Now'] = np.where(agg_df['Impr_Now'] > 0, agg_df['Clicks_Now'] / agg_df['Impr_Now'] * 100, 0)
    
    # Calculate changes
    agg_df['Clicks_Delta'] = agg_df['Clicks_Now'] - agg_df['Clicks_Prev']
    agg_df['Impr_Delta'] = agg_df['Impr_Now'] - agg_df['Impr_Prev']
    agg_df['Clicks_Pct_Change'] = np.where(agg_df['Clicks_Prev'] > 0, 
                                           agg_df['Clicks_Delta'] / agg_df['Clicks_Prev'] * 100, 0)
    
    # Sort by current clicks
    agg_df = agg_df.sort_values('Clicks_Now', ascending=False).reset_index(drop=True)
    
    # Top performers
    top_pages_by_clicks = agg_df.head(25)
    top_ctr_pages = agg_df[agg_df['CTR_Now'] > 0].sort_values('CTR_Now', ascending=False).head(25)
    biggest_gainers = agg_df.sort_values('Clicks_Delta', ascending=False).head(20)
    biggest_losers = agg_df.sort_values('Clicks_Delta', ascending=True).head(20)
    
    # Countries analysis
    countries_analysis = None
    if countries_df is not None and not countries_df.empty:
        try:
            countries_analysis = analyze_countries_data(countries_df)
        except:
            countries_analysis = None
    
    return {
        'total_pages': len(agg_df),  # Now count distinct normalized URLs
        'total_clicks_now': agg_df['Clicks_Now'].sum(),
        'total_clicks_prev': agg_df['Clicks_Prev'].sum(),
        'total_clicks_delta': agg_df['Clicks_Delta'].sum(),
        'avg_ctr': agg_df['CTR_Now'].mean(),
        'top_pages_by_clicks': top_pages_by_clicks,
        'top_ctr_pages': top_ctr_pages,
        'biggest_gainers': biggest_gainers,
        'biggest_losers': biggest_losers,
        'countries_analysis': countries_analysis,
        'raw_data': agg_df
    }

def analyze_countries_data(countries_df):
    """Analyze countries data for mapping"""
    
    # Find columns
    country_col = find_column(countries_df.columns, ['country', 'countries'])
    clicks_col = find_column(countries_df.columns, ['clicks'])
    impr_col = find_column(countries_df.columns, ['impressions'])
    
    if not country_col or not clicks_col:
        return None
    
    # Build countries dataframe
    countries_work = pd.DataFrame()
    countries_work['Country'] = countries_df[country_col].astype(str).str.strip()
    countries_work['Clicks'] = pd.to_numeric(countries_df[clicks_col], errors='coerce')
    
    if impr_col:
        countries_work['Impressions'] = pd.to_numeric(countries_df[impr_col], errors='coerce')
        countries_work['CTR'] = np.where(countries_work['Impressions'] > 0, 
                                        countries_work['Clicks'] / countries_work['Impressions'] * 100, 0)
    
    # Clean and sort
    countries_work = countries_work[countries_work['Country'].notna() & (countries_work['Clicks'] > 0)].copy()
    countries_work = countries_work.sort_values('Clicks', ascending=False)
    
    return countries_work

def display_gsc_results(results):
    """Display GSC page performance results"""
    
    # Key metrics
    st.markdown('<div class="section-header">📈 GSC Page Performance Summary</div>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            label="Total Pages",
            value=f"{results['total_pages']:,}"
        )
    
    with col2:
        st.metric(
            label="Total Clicks Change",
            value=f"{results['total_clicks_delta']:,}",
            delta=f"{((results['total_clicks_delta'] / results['total_clicks_prev']) * 100) if results['total_clicks_prev'] > 0 else 0:+.1f}%",
            delta_color="normal"  # Always use normal: green=good/increase, red=bad/decrease
        )
    
    with col3:
        st.metric(
            label="Current Total Clicks",
            value=f"{results['total_clicks_now']:,}"
        )
    
    with col4:
        st.metric(
            label="Average CTR",
            value=f"{results['avg_ctr']:.2f}%",
            help="Average click-through rate across all pages"
        )
    
    # Top Pages Analysis
    st.markdown('<div class="section-header">🏆 Top Performing Pages</div>', unsafe_allow_html=True)
    
    # Use tabs for different views
    clicks_tab, ctr_tab, gainers_tab, losers_tab = st.tabs(["📊 By Clicks", "🎯 By CTR", "📈 Biggest Gainers", "📉 Biggest Losers"])
    
    with clicks_tab:
        st.markdown("**Pages with highest current click volume**")
        if results.get('top_pages_by_clicks') is not None and not results['top_pages_by_clicks'].empty:
            display_cols = ['Page', 'Clicks_Now', 'Impr_Now', 'CTR_Now']
            if 'Position' in results['top_pages_by_clicks'].columns:
                display_cols.append('Position')
            
            display_df = results['top_pages_by_clicks'][display_cols].copy()
            display_df.columns = ['Page', 'Clicks', 'Impressions', 'CTR %'] + (['Avg Position'] if len(display_cols) > 4 else [])
            st.dataframe(display_df, use_container_width=True, hide_index=True, height=400)
        else:
            st.info("No click data available")
    
    with ctr_tab:
        st.markdown("**Pages with highest click-through rates**")
        if results.get('top_ctr_pages') is not None and not results['top_ctr_pages'].empty:
            display_cols = ['Page', 'Clicks_Now', 'CTR_Now']
            if 'Position' in results['top_ctr_pages'].columns:
                display_cols.append('Position')
            
            display_df = results['top_ctr_pages'][display_cols].copy()
            display_df.columns = ['Page', 'Clicks', 'CTR %'] + (['Avg Position'] if len(display_cols) > 3 else [])
            st.dataframe(display_df, use_container_width=True, hide_index=True, height=400)
        else:
            st.info("No CTR data available")
    
    with gainers_tab:
        st.markdown("**Pages with biggest click increases**")
        if results.get('biggest_gainers') is not None and not results['biggest_gainers'].empty:
            display_cols = ['Page', 'Clicks_Prev', 'Clicks_Now', 'Clicks_Delta', 'Clicks_Pct_Change']
            display_df = results['biggest_gainers'][display_cols].copy()
            display_df.columns = ['Page', 'Previous Clicks', 'Current Clicks', 'Clicks Δ', 'Change %']
            st.dataframe(display_df, use_container_width=True, hide_index=True, height=400)
        else:
            st.info("No gainers data available")
    
    with losers_tab:
        st.markdown("**Pages with biggest click decreases**")
        if results.get('biggest_losers') is not None and not results['biggest_losers'].empty:
            display_cols = ['Page', 'Clicks_Prev', 'Clicks_Now', 'Clicks_Delta', 'Clicks_Pct_Change']
            display_df = results['biggest_losers'][display_cols].copy()
            display_df.columns = ['Page', 'Previous Clicks', 'Current Clicks', 'Clicks Δ', 'Change %']
            st.dataframe(display_df, use_container_width=True, hide_index=True, height=400)
        else:
            st.info("No losers data available")
    
    # Countries Map (if available)
    if results['countries_analysis'] is not None and not results['countries_analysis'].empty:
        st.markdown('<div class="section-header">🗺️ Geographic Performance</div>', unsafe_allow_html=True)
        
        countries_data = results['countries_analysis']
        
        # Create world map
        fig_map = go.Figure(data=go.Choropleth(
            locations=countries_data['Country'],
            z=countries_data['Clicks'],
            locationmode='country names',
            colorscale='Blues',
            autocolorscale=False,
            text=countries_data['Country'],
            marker_line_color='darkgray',
            marker_line_width=0.5,
            colorbar_title="Clicks"
        ))
        
        fig_map.update_layout(
            title=dict(text='Click Distribution by Country', font=dict(size=20)),
            geo=dict(
                showframe=False,
                showcoastlines=True,
                projection_type='equirectangular'
            ),
            height=500,
            margin=dict(l=20, r=20, t=60, b=20)
        )
        
        st.plotly_chart(fig_map, use_container_width=True, config={'displayModeBar': False})
        
        # Countries table
        st.markdown("**Top Countries by Clicks**")
        countries_display = countries_data.head(15).copy()
        if 'CTR' in countries_display.columns:
            countries_display['CTR'] = countries_display['CTR'].round(2)
        st.dataframe(countries_display, use_container_width=True, hide_index=True, height=300)
    
    # Download section
    st.markdown('<div class="section-header">📥 Download GSC Results</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        gsc_report = create_gsc_summary_report(results)
        st.download_button(
            label="📄 Download GSC Analysis Report",
            data=gsc_report,
            file_name=f"gsc_page_analysis_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain"
        )
    
    with col2:
        csv_buffer = io.StringIO()
        results['raw_data'].to_csv(csv_buffer, index=False)
        
        st.download_button(
            label="📊 Download GSC Data (CSV)",
            data=csv_buffer.getvalue(),
            file_name=f"gsc_page_data_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv"
        )

def create_gsc_summary_report(results):
    """Create downloadable GSC analysis report"""
    
    clicks_change_pct = (results['total_clicks_delta'] / results['total_clicks_prev'] * 100) if results['total_clicks_prev'] > 0 else 0
    
    report = f"""
GSC PAGE PERFORMANCE ANALYSIS REPORT
Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

===========================================
EXECUTIVE SUMMARY
===========================================

Total Pages Analyzed: {results['total_pages']:,}
Total Clicks Change: {results['total_clicks_delta']:,} ({clicks_change_pct:+.1f}%)
Current Total Clicks: {results['total_clicks_now']:,}
Average CTR: {results['avg_ctr']:.2f}%

===========================================
TOP PAGES BY CLICKS
===========================================

"""
    
    for _, row in results['top_pages_by_clicks'].head(15).iterrows():
        position_info = f" | Avg Pos: {row['Position']:.1f}" if 'Position' in row and pd.notna(row['Position']) else ""
        report += f"• {row['Page']} | {row['Clicks_Now']:.0f} clicks | CTR: {row['CTR_Now']:.2f}%{position_info}\n"
    
    report += f"""

===========================================
BIGGEST CLICK GAINERS
===========================================

"""
    
    for _, row in results['biggest_gainers'].head(10).iterrows():
        report += f"• {row['Page']} | +{row['Clicks_Delta']:.0f} clicks ({row['Clicks_Pct_Change']:+.1f}%)\n"
    
    if results['countries_analysis'] is not None and not results['countries_analysis'].empty:
        report += f"""

===========================================
TOP COUNTRIES BY CLICKS
===========================================

"""
        
        for _, row in results['countries_analysis'].head(10).iterrows():
            ctr_info = f" | CTR: {row['CTR']:.2f}%" if 'CTR' in row and pd.notna(row['CTR']) else ""
            report += f"• {row['Country']} | {row['Clicks']:.0f} clicks{ctr_info}\n"
    
    report += f"""

===========================================
"""
    
    return report

def competitor_analysis():
    """Analyze competitor rankings and gaps"""
    st.markdown('<div class="section-header">🏁 Competitor Gap Analysis</div>', unsafe_allow_html=True)
    
    # Modern instruction design
    with st.container():
        st.markdown("### 📊 Analysis Overview")
        
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown("""
            This analysis examines your competitive landscape to understand:
            
            **🎯 Key Questions Answered:**
            - Who are your real search competitors (not just business rivals)?
            - Where do competitors consistently outrank you?
            - Which of your declining keywords show competitive pressure?
            - What specific keyword gaps present opportunities?
            """)
        
        with col2:
            st.info("""
            **💡 Strategic Value**
            
            Identifies the domains shaping your SERPs and reveals specific keyword opportunities to target.
            """)
    
    # File requirements in expandable section
    with st.expander("📁 **File Requirements & Setup**", expanded=False):
        st.markdown("""
        **Required Files:** 2 files minimum, 3+ for detailed analysis
        
        | File | Purpose | Export From |
        |------|---------|-------------|
        | **Semrush Competitors** | Identify top competitors | Organic Research → Competitors |
        | **Your Positions (current)** | Your current rankings | Organic Research → Positions |
        | **Competitor Positions** | Optional: Detailed gaps | Individual competitor Position exports |
        
        **📋 Export Steps:**
        1. **Competitors**: Export from Competitors tab (shows relevance & overlap)
        2. **Your Positions**: Current month positions export  
        3. **Optional**: Export positions for top 3-5 competitors individually
        
        **🔍 Analysis Method:**
        - Identifies top competitors by relevance/keyword overlap
        - Counts where each competitor outranks you
        - Shows specific keyword gap opportunities
        - Focuses on competitive pressure for declining queries
        """)
    
    # Key insights preview
    st.markdown("### 🎯 Analysis Insights You'll Get")
    
    insight_col1, insight_col2, insight_col3, insight_col4 = st.columns(4)
    
    with insight_col1:
        st.markdown("""
        **🥇 Top Competitors**
        - Real search competitors by relevance
        - Keyword overlap analysis
        """)
    
    with insight_col2:
        st.markdown("""
        **📊 Outrank Counts**
        - Where competitors beat you
        - Win/loss ratios by competitor
        """)
    
    with insight_col3:
        st.markdown("""
        **🎯 Gap Opportunities**
        - Specific keywords to target
        - Competitive displacement analysis
        """)
    
    with insight_col4:
        st.markdown("""
        **📉 Pressure Analysis**
        - Competitors affecting declining queries
        - Strategic counter-moves
        """)
    
    st.markdown("---")
    
    # File upload section
    st.markdown("### 📤 Upload Your Data Files")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 🏆 Semrush Competitors (Required)")
        competitors_file = st.file_uploader(
            "Upload Semrush Competitors file",
            type=['csv', 'xlsx', 'xls'],
            key="competitors_file",
            help="Export from Semrush: Organic Research → Competitors"
        )
        
    with col2:
        st.markdown("#### 📊 Your Positions (Required)")
        your_positions_file = st.file_uploader(
            "Upload Your Semrush Positions file",
            type=['csv', 'xlsx', 'xls'],
            key="your_positions_file",
            help="Export from Semrush: Organic Research → Positions (current)"
        )
    
    # Optional competitor positions
    st.markdown("#### 🎯 Competitor Positions (Optional - for detailed gap analysis)")
    competitor_positions_files = st.file_uploader(
        "Upload competitor position files (one per competitor)",
        type=['csv', 'xlsx', 'xls'],
        accept_multiple_files=True,
        key="competitor_positions_files",
        help="Optional: Individual position exports for top competitors"
    )
    
    # Process files if minimum required files uploaded
    if competitors_file is not None and your_positions_file is not None:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            run_competitor_analysis = st.button("🚀 Run Competitor Analysis", key="run_competitors", type="primary", use_container_width=True)
        
        if run_competitor_analysis:
            with st.spinner("🔄 Analyzing competitive landscape..."):
                try:
                    # Load main files
                    competitors_df = normalize_columns(read_uploaded_file(competitors_file))
                    your_positions_df = normalize_columns(read_uploaded_file(your_positions_file))
                    
                    # Load optional competitor files
                    competitor_data = {}
                    if competitor_positions_files:
                        for comp_file in competitor_positions_files:
                            try:
                                comp_df = normalize_columns(read_uploaded_file(comp_file))
                                # Try to infer domain from filename or data
                                domain = infer_competitor_domain(comp_file.name, comp_df)
                                if domain:
                                    competitor_data[domain] = comp_df
                            except Exception as e:
                                st.warning(f"Could not process {comp_file.name}: {str(e)}")
                    
                    # Validate data
                    validation_passed, validation_message = validate_competitor_data(competitors_df, your_positions_df)
                    
                    if not validation_passed:
                        st.error(validation_message)
                        st.stop()
                    
                    # Perform analysis
                    competitor_results = analyze_competitor_gaps(competitors_df, your_positions_df, competitor_data)
                    
                    # Display results
                    display_competitor_results(competitor_results)
                    
                except Exception as e:
                    st.error(f"❌ Error processing files: {str(e)}")
                    st.info("💡 Please ensure you've uploaded valid Semrush files")
    else:
        missing = []
        if competitors_file is None:
            missing.append("Competitors file")
        if your_positions_file is None:
            missing.append("Your Positions file")
        st.info(f"📤 Please upload: {', '.join(missing)}")


def analyze_page_performance(df):
    """Analyze page performance patterns following the prototype methodology"""
    
    # Find and rename columns
    url_col = find_column(df.columns, ['url', 'page', 'landing page'])
    
    # Look for Traffic column (prefer exact match)
    traffic_exact = [c for c in df.columns if c.lower().strip() == 'traffic']
    traffic_col = traffic_exact[0] if traffic_exact else find_column(df.columns, ['traffic'])
    
    traffic_pct_col = find_column(df.columns, ['traffic %', 'traffic%', 'traffic (%)', 'share'])
    keywords_col = find_column(df.columns, ['number of keywords', 'keywords', 'num. keywords', 'kws'])
    
    # Build working dataframe
    work_df = pd.DataFrame()
    work_df['URL'] = df[url_col].astype(str).str.strip()
    
    # Normalize URLs to handle trailing slashes and case variations
    def normalize_url(url):
        if pd.isna(url) or url == '':
            return url
        url = str(url).strip()
        # Remove trailing slash (but keep if it's just the domain)
        if url.endswith('/') and url.count('/') > 2:
            url = url.rstrip('/')
        # Convert to lowercase for consistent comparison
        return url.lower()
    
    work_df['URL_Normalized'] = work_df['URL'].apply(normalize_url)
    work_df['Traffic'] = pd.to_numeric(df[traffic_col].astype(str).str.replace(',', ''), errors='coerce')
    
    if traffic_pct_col:
        # Handle percentage formats
        traffic_pct_series = df[traffic_pct_col]
        if traffic_pct_series.astype(str).str.contains('%').any():
            work_df['Traffic_Pct'] = pd.to_numeric(traffic_pct_series.astype(str).str.replace('%', ''), errors='coerce')
        else:
            pct_numeric = pd.to_numeric(traffic_pct_series, errors='coerce')
            # If values are between 0-1, assume they're fractions, convert to percentages
            if pct_numeric.max() <= 1.0:
                work_df['Traffic_Pct'] = pct_numeric * 100
            else:
                work_df['Traffic_Pct'] = pct_numeric
    else:
        # Calculate traffic percentage if not provided
        total_traffic = work_df['Traffic'].sum()
        work_df['Traffic_Pct'] = (work_df['Traffic'] / total_traffic * 100) if total_traffic > 0 else 0
    
    if keywords_col:
        work_df['Keywords'] = pd.to_numeric(df[keywords_col].astype(str).str.replace(',', ''), errors='coerce').fillna(0).astype(int)
    else:
        work_df['Keywords'] = np.nan
    
    # Clean data
    work_df = work_df[work_df['URL'].notna() & work_df['URL'].ne('') & work_df['Traffic'].notna()].copy()
    
    # Aggregate duplicates by normalized URL (sum traffic, keep first URL for display)
    agg_df = work_df.groupby('URL_Normalized').agg({
        'URL': 'first',  # Keep first URL for display
        'Traffic': 'sum',  # Sum traffic for duplicates
        'Traffic_Pct': 'sum',  # Sum traffic percentage
        'Keywords': 'sum' if 'Keywords' in work_df.columns else 'first'
    }).reset_index()
    
    # Recalculate traffic percentage after aggregation
    total_traffic = agg_df['Traffic'].sum()
    agg_df['Traffic_Pct'] = (agg_df['Traffic'] / total_traffic * 100) if total_traffic > 0 else 0
    
    # Sort by traffic
    agg_df = agg_df.sort_values('Traffic', ascending=False).reset_index(drop=True)
    
    # 1. Pareto Analysis
    agg_df['Cumulative_Pct'] = agg_df['Traffic_Pct'].cumsum().clip(upper=100)
    
    def pages_to_threshold(threshold):
        if agg_df['Cumulative_Pct'].empty:
            return np.nan
        idx = np.argmax(agg_df['Cumulative_Pct'].values >= threshold)
        return int(idx + 1) if agg_df['Cumulative_Pct'].iloc[-1] >= threshold else len(agg_df)
    
    pareto_thresholds = {
        '50%': pages_to_threshold(50),
        '80%': pages_to_threshold(80), 
        '90%': pages_to_threshold(90)
    }
    
    # 2. Efficiency Analysis (Traffic per Keyword)
    efficiency_df = pd.DataFrame()
    if not agg_df['Keywords'].isna().all():
        eff_df = agg_df[agg_df['Keywords'] > 0].copy()
        min_keywords = max(5, int(np.median(eff_df['Keywords']))) if len(eff_df) > 0 else 5
        min_keywords = min(min_keywords, 20)  # Cap at 20 for broader analysis
        
        eff_df = eff_df[eff_df['Keywords'] >= min_keywords].copy()
        eff_df['TPK'] = (eff_df['Traffic'] / eff_df['Keywords']).round(2)
        efficiency_df = eff_df.sort_values('TPK', ascending=False)
    
    # 3. Directory Analysis
    def extract_first_directory(url):
        try:
            from urllib.parse import urlparse
            path = urlparse(url).path.strip('/')
            return ('/' + path.split('/')[0]) if path else '/'
        except:
            return '/'
    
    agg_df['Directory'] = agg_df['URL'].apply(extract_first_directory)
    
    directory_analysis = (agg_df.groupby('Directory')
                         .agg({
                             'Traffic': 'sum',
                             'URL': 'count',  # Page count
                             'Keywords': 'sum'
                         })
                         .rename(columns={'URL': 'Pages'})
                         .sort_values('Traffic', ascending=False)
                         .reset_index())
    
    directory_analysis['Traffic_Pct'] = (directory_analysis['Traffic'] / directory_analysis['Traffic'].sum() * 100).round(2)
    directory_analysis['Avg_Traffic_Per_Page'] = (directory_analysis['Traffic'] / directory_analysis['Pages']).round(2)
    
    # 4. Long-tail Opportunities
    longtail_df = pd.DataFrame()
    if not agg_df['Keywords'].isna().all():
        opp_df = agg_df[agg_df['Keywords'] > 0].copy()
        opp_df['TPK'] = (opp_df['Traffic'] / opp_df['Keywords']).replace([np.inf, -np.inf], np.nan)
        
        # High breadth (75th percentile keywords), low efficiency (25th percentile TPK)
        kw_threshold = opp_df['Keywords'].quantile(0.75)
        tpk_threshold = opp_df['TPK'].quantile(0.25)
        
        longtail_df = opp_df[(opp_df['Keywords'] >= kw_threshold) & (opp_df['TPK'] <= tpk_threshold)].copy()
        longtail_df = longtail_df.sort_values(['Keywords', 'TPK'], ascending=[False, True])
    
    return {
        'total_pages': len(agg_df),  # Now count distinct normalized URLs
        'total_traffic': agg_df['Traffic'].sum(),
        'pareto_data': agg_df[['URL', 'Traffic', 'Traffic_Pct', 'Cumulative_Pct']].copy(),
        'pareto_thresholds': pareto_thresholds,
        'top_pages': agg_df.head(25),
        'efficiency_analysis': efficiency_df.head(25) if not efficiency_df.empty else pd.DataFrame(),
        'directory_analysis': directory_analysis.head(15),
        'longtail_opportunities': longtail_df.head(25) if not longtail_df.empty else pd.DataFrame(),
        'raw_data': agg_df
    }

def display_pages_results(results):
    """Display page performance analysis results"""
    
    # Key metrics
    st.markdown('<div class="section-header">📈 Page Performance Summary</div>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            label="Total Pages",
            value=f"{results['total_pages']:,}"
        )
    
    with col2:
        st.metric(
            label="Total Traffic",
            value=f"{results['total_traffic']:,.0f}",
            help="Estimated monthly organic traffic"
        )
    
    with col3:
        # Calculate top 10 pages traffic share - more actionable metric
        top_10_traffic = results['top_pages'].head(10)['Traffic'].sum() if len(results['top_pages']) >= 10 else results['top_pages']['Traffic'].sum()
        top_10_share = (top_10_traffic / results['total_traffic'] * 100) if results['total_traffic'] > 0 else 0
        
        st.metric(
            label="Top 10 Pages Share",
            value=f"{top_10_share:.1f}%",
            help="Percentage of total traffic from your top 10 pages"
        )
    
    with col4:
        if not results['efficiency_analysis'].empty:
            avg_tpk = results['efficiency_analysis']['TPK'].mean()
            st.metric(
                label="Avg TPK (Top Pages)",
                value=f"{avg_tpk:.1f}",
                help="Average Traffic per Keyword for efficient pages"
            )
        else:
            st.metric(label="TPK Analysis", value="N/A", help="Keywords data not available")
    
    # Pareto Analysis
    st.markdown('<div class="section-header">📊 Traffic Concentration (Pareto Analysis)</div>', unsafe_allow_html=True)
    
    # Pareto curve chart
    pareto_data = results['pareto_data']
    fig_pareto = go.Figure()
    
    fig_pareto.add_trace(go.Scatter(
        x=list(range(1, len(pareto_data) + 1)),
        y=pareto_data['Cumulative_Pct'].values,
        mode='lines+markers',
        name='Cumulative Traffic %',
        line=dict(color='#3498db', width=3),
        marker=dict(size=4)
    ))
    
    # Add threshold lines
    for threshold, pages_needed in results['pareto_thresholds'].items():
        if pages_needed and not np.isnan(pages_needed):
            fig_pareto.add_hline(y=int(threshold.replace('%', '')), 
                                line_dash="dash", 
                                line_color="red",
                                annotation_text=f"{threshold} - {pages_needed} pages")
    
    fig_pareto.update_layout(
        title=dict(text='Traffic Concentration Curve (Pareto Analysis)', font=dict(size=20)),
        xaxis_title='Pages (Ranked by Traffic)',
        yaxis_title='Cumulative Traffic Share (%)',
        height=500,
        margin=dict(l=60, r=60, t=80, b=60),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        showlegend=False
    )
    
    st.plotly_chart(fig_pareto, use_container_width=True, config={'displayModeBar': False})
    
    # Pareto summary table
    pareto_summary = pd.DataFrame({
        'Traffic Threshold': ['50% of traffic', '80% of traffic', '90% of traffic'],
        'Pages Needed': [results['pareto_thresholds']['50%'], 
                        results['pareto_thresholds']['80%'], 
                        results['pareto_thresholds']['90%']]
    })
    st.dataframe(pareto_summary, use_container_width=True, hide_index=True)
    
    # Top Pages Analysis
    st.markdown('<div class="section-header">🏆 Top Performing Pages</div>', unsafe_allow_html=True)
    
    # Show top pages with key metrics
    top_pages_display = results['top_pages'][['URL', 'Traffic', 'Traffic_Pct']].copy()
    if 'Keywords' in results['top_pages'].columns:
        top_pages_display['Keywords'] = results['top_pages']['Keywords']
    
    top_pages_display.columns = ['URL', 'Traffic', 'Traffic %'] + (['Keywords'] if 'Keywords' in top_pages_display.columns else [])
    st.dataframe(top_pages_display, use_container_width=True, hide_index=True, height=400)
    
    # Efficiency and Directory Analysis
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown('<div class="section-header">⚡ Efficiency Leaders (Traffic per Keyword)</div>', unsafe_allow_html=True)
        
        if not results['efficiency_analysis'].empty:
            eff_display = results['efficiency_analysis'][['URL', 'Traffic', 'Keywords', 'TPK']].copy()
            eff_display.columns = ['URL', 'Traffic', 'Keywords', 'TPK']
            st.dataframe(eff_display, use_container_width=True, hide_index=True, height=350)
            st.caption("*Pages with highest traffic per keyword - strong intent match signals*")
        else:
            st.info("Keywords data not available for efficiency analysis")
    
    with col2:
        st.markdown('<div class="section-header">🗂️ Directory Performance</div>', unsafe_allow_html=True)
        
        dir_display = results['directory_analysis'][['Directory', 'Pages', 'Traffic', 'Traffic_Pct', 'Avg_Traffic_Per_Page']].copy()
        dir_display.columns = ['Directory', 'Pages', 'Traffic', 'Traffic %', 'Avg/Page']
        st.dataframe(dir_display, use_container_width=True, hide_index=True, height=350)
        st.caption("*Traffic distribution by content directory/hub*")
    
    # Directory visualization
    st.markdown('<div class="section-header">📊 Directory Traffic Distribution</div>', unsafe_allow_html=True)
    
    top_dirs = results['directory_analysis'].head(12)
    fig_dirs = go.Figure(data=[
        go.Bar(
            y=top_dirs['Directory'],
            x=top_dirs['Traffic'],
            orientation='h',
            marker_color='#3498db',
            text=[f"{val:,.0f}" for val in top_dirs['Traffic']],
            textposition='outside'
        )
    ])
    
    fig_dirs.update_layout(
        title=dict(text='Top Directories by Traffic', font=dict(size=20)),
        xaxis_title='Traffic',
        yaxis_title='Directory',
        height=500,
        margin=dict(l=150, r=60, t=80, b=60),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        yaxis=dict(autorange='reversed')  # Biggest at top
    )
    
    st.plotly_chart(fig_dirs, use_container_width=True, config={'displayModeBar': False})
    
    # Long-tail Opportunities
    if not results['longtail_opportunities'].empty:
        st.markdown('<div class="section-header">🎯 Long-tail Optimization Opportunities</div>', unsafe_allow_html=True)
        st.markdown("*Pages with many keywords but low traffic per keyword - candidates for internal links, schema, and intent expansion*")
        
        longtail_display = results['longtail_opportunities'][['URL', 'Traffic', 'Keywords', 'TPK']].copy()
        longtail_display.columns = ['URL', 'Traffic', 'Keywords', 'TPK']
        st.dataframe(longtail_display, use_container_width=True, hide_index=True, height=400)
    
    # Strategic insights - AI-powered
    st.markdown('<div class="section-header">💡 Strategic Insights</div>', unsafe_allow_html=True)
    
    st.info("💡 **AI-Powered Insights:** Strategic insights are generated using AI. Set your OpenAI API key as an environment variable `OPENAI_API_KEY` or in Streamlit secrets.")
    
    with st.spinner("🤖 Generating AI-powered strategic insights..."):
        analysis_summary = create_pages_analysis_summary(results)
        insights = generate_chatgpt_insights(analysis_summary, "pages")
        st.markdown(f'<div class="insight-box">{insights.replace(chr(10), "<br>")}</div>', unsafe_allow_html=True)
    
    # Download section
    st.markdown('<div class="section-header">📥 Download Results</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        summary_report = create_pages_summary_report(results)
        st.download_button(
            label="📄 Download Page Analysis Report",
            data=summary_report,
            file_name=f"page_performance_analysis_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain"
        )
    
    with col2:
        # Convert analysis to CSV
        csv_buffer = io.StringIO()
        results['raw_data'].to_csv(csv_buffer, index=False)
        
        st.download_button(
            label="📊 Download Full Data (CSV)",
            data=csv_buffer.getvalue(),
            file_name=f"page_performance_data_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv"
        )

def generate_pages_insights(results):
    """Generate strategic insights from page performance analysis"""
    insights = []
    
    # Traffic concentration analysis
    pages_50 = results['pareto_thresholds']['50%']
    pages_80 = results['pareto_thresholds']['80%']
    total_pages = results['total_pages']
    
    concentration_50 = (pages_50 / total_pages * 100) if pages_50 and total_pages > 0 else 0
    
    if concentration_50 < 5:
        insights.append(f"<b>🟢 Excellent Traffic Distribution:</b> Only {pages_50} pages ({concentration_50:.1f}%) drive 50% of traffic - very concentrated and efficient.")
    elif concentration_50 < 15:
        insights.append(f"<b>🟢 Good Concentration:</b> {pages_50} pages ({concentration_50:.1f}%) drive 50% of traffic - healthy focus on high-value content.")
    else:
        insights.append(f"<b>🟡 Distributed Traffic:</b> {pages_50} pages ({concentration_50:.1f}%) needed for 50% of traffic - consider strengthening top performers.")
    
    # Efficiency insights
    if not results['efficiency_analysis'].empty:
        top_tpk = results['efficiency_analysis']['TPK'].iloc[0] if len(results['efficiency_analysis']) > 0 else 0
        insights.append(f"<b>⚡ Top Efficiency:</b> Your most efficient page generates {top_tpk:.1f} traffic per keyword - analyze and replicate this content pattern.")
    
    # Directory insights
    if not results['directory_analysis'].empty:
        top_dir = results['directory_analysis'].iloc[0]
        dir_concentration = top_dir['Traffic_Pct']
        insights.append(f"<b>🗂️ Content Hub Leader:</b> The '{top_dir['Directory']}' directory drives {dir_concentration:.1f}% of traffic from {top_dir['Pages']} pages.")
    
    # Long-tail insights
    if not results['longtail_opportunities'].empty:
        longtail_count = len(results['longtail_opportunities'])
        insights.append(f"<b>🎯 Optimization Potential:</b> {longtail_count} pages have high keyword counts but low efficiency - prime candidates for internal linking and content enhancement.")
    
    # Risk assessment
    if concentration_50 < 10:
        insights.append("<b>⚠️ Concentration Risk:</b> Heavy reliance on few pages - diversify traffic sources and defend top performers against competitive threats.")
    
    return "<br><br>".join(insights)

def create_pages_summary_report(results):
    """Create downloadable page performance report"""
    
    report = f"""
PAGE PERFORMANCE ANALYSIS REPORT
Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

===========================================
EXECUTIVE SUMMARY
===========================================

Total Pages Analyzed: {results['total_pages']:,}
Total Estimated Traffic: {results['total_traffic']:,.0f}

Traffic Concentration (Pareto Analysis):
• 50% of traffic: {results['pareto_thresholds']['50%']} pages
• 80% of traffic: {results['pareto_thresholds']['80%']} pages  
• 90% of traffic: {results['pareto_thresholds']['90%']} pages

===========================================
TOP PERFORMING PAGES (Traffic Leaders)
===========================================

"""
    
    for _, row in results['top_pages'].head(15).iterrows():
        keywords_info = f" | {row['Keywords']:.0f} keywords" if 'Keywords' in row and pd.notna(row['Keywords']) else ""
        report += f"• {row['URL']} | {row['Traffic']:.0f} traffic ({row['Traffic_Pct']:.1f}%){keywords_info}\n"
    
    if not results['efficiency_analysis'].empty:
        report += f"""

===========================================
EFFICIENCY LEADERS (Traffic per Keyword)
===========================================

"""
        for _, row in results['efficiency_analysis'].head(10).iterrows():
            report += f"• {row['URL']} | TPK: {row['TPK']:.1f} ({row['Traffic']:.0f} traffic / {row['Keywords']:.0f} keywords)\n"
    
    report += f"""

===========================================
CONTENT HUB ANALYSIS (Directory Performance)  
===========================================

"""
    
    for _, row in results['directory_analysis'].head(10).iterrows():
        report += f"• {row['Directory']} | {row['Traffic']:.0f} traffic ({row['Traffic_Pct']:.1f}%) from {row['Pages']} pages\n"
    
    if not results['longtail_opportunities'].empty:
        report += f"""

===========================================
LONG-TAIL OPTIMIZATION OPPORTUNITIES
===========================================

"""
        for _, row in results['longtail_opportunities'].head(10).iterrows():
            report += f"• {row['URL']} | {row['Keywords']:.0f} keywords, TPK: {row['TPK']:.1f}\n"
    
    report += f"""

===========================================
STRATEGIC INSIGHTS
===========================================

{generate_pages_insights(results).replace('<b>', '').replace('</b>', '').replace('<br><br>', '\n\n').replace('🟢', '• ').replace('🟡', '• ').replace('⚡', '• ').replace('🗂️', '• ').replace('🎯', '• ').replace('⚠️', '• ')}

===========================================
"""
    
    return report

def query_gains_losses_analysis():
    """Analyze query-level gains and losses from GSC"""
    st.markdown('<div class="section-header">🎯 Query Performance Analysis</div>', unsafe_allow_html=True)
    
    # Modern instruction design
    with st.container():
        st.markdown("### 📊 Analysis Overview")
        
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown("""
            This analysis examines search query performance to identify:
            
            **🎯 Key Questions Answered:**
            - Which search terms are driving the most additional clicks?
            - Which queries are losing traffic and need attention?
            - Are click losses due to ranking drops or CTR pressure?
            - Which queries show SERP feature impact?
            """)
        
        with col2:
            st.info("""
            **💡 Strategic Value**
            
            Separates demand growth from execution issues to prioritize the right optimization efforts.
            """)
    
    # File requirements in expandable section
    with st.expander("📁 **File Requirements & Setup**", expanded=False):
        st.markdown("""
        **Required Files:** 1-2 files for comprehensive analysis
        
        | File | Purpose | Export From |
        |------|---------|-------------|
        | **GSC Queries Compare** | Primary analysis | Search Console → Search Results → Queries (Compare view) |
        | **Semrush Positions** | Optional enrichment | Organic Research → Positions (current period) |
        
        **📋 GSC Export Steps:**
        1. Go to Search Results in Google Search Console
        2. Set: Search type = Web
        3. Set: Date → Compare → Last 3 months vs Same period last year
        4. Click Queries tab
        5. Export → CSV or Excel
        
        **🔍 Analysis Method:**
        - Clicks Δ = Current - Previous period
        - Impressions Δ separates demand vs execution issues
        - Position context from Semrush (if provided)
        """)
    
    # Key insights preview
    st.markdown("### 🎯 Analysis Insights You'll Get")
    
    insight_col1, insight_col2, insight_col3, insight_col4 = st.columns(4)
    
    with insight_col1:
        st.markdown("""
        **📈 Top Winners**
        - Queries with biggest click gains
        - Growth opportunity patterns
        """)
    
    with insight_col2:
        st.markdown("""
        **📉 Top Losers**
        - Queries losing the most clicks
        - Decline root cause analysis
        """)
    
    with insight_col3:
        st.markdown("""
        **🎯 CTR Pressure**
        - Impressions up, clicks down cases
        - SERP feature impact identification
        """)
    
    with insight_col4:
        st.markdown("""
        **🏆 Ranking Wins**
        - Clicks up with flat impressions
        - Position improvement validation
        """)
    
    st.markdown("---")
    
    # File upload section
    st.markdown("### 📤 Upload Your Data Files")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 📊 GSC Queries Compare (Required)")
        gsc_queries_file = st.file_uploader(
            "Upload GSC Queries Compare file",
            type=['csv', 'xlsx', 'xls'],
            key="gsc_queries",
            help="Export from GSC: Search Results → Queries (Compare view)"
        )
        
    with col2:
        st.markdown("#### 📊 Semrush Positions (Optional)")
        semrush_positions_file = st.file_uploader(
            "Upload Semrush Positions file (for enrichment)",
            type=['csv', 'xlsx', 'xls'],
            key="semrush_positions_enrich",
            help="Optional: Adds current position context to query analysis"
        )
    
    # Process files if main file uploaded
    if gsc_queries_file is not None:
        # Add Run Analysis button (centered)
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            run_query_analysis = st.button("🚀 Run Query Analysis", key="run_queries", type="primary", use_container_width=True)
        
        # Display results outside column context for full width
        if run_query_analysis:
            with st.spinner("🔄 Analyzing query performance..."):
                try:
                    # Load and validate data
                    gsc_df = normalize_columns(read_uploaded_file(gsc_queries_file))
                    
                    # Optional Semrush enrichment
                    semrush_df = None
                    if semrush_positions_file is not None:
                        semrush_df = normalize_columns(read_uploaded_file(semrush_positions_file))
                    
                    # Validate required columns
                    validation_passed, validation_message = validate_query_data(gsc_df)
                    
                    if not validation_passed:
                        st.error(validation_message)
                        st.stop()
                    
                    # Perform analysis
                    query_results = analyze_query_performance(gsc_df, semrush_df)
                    
                    # Display results - FULL WIDTH
                    display_query_results(query_results)
                    
                except Exception as e:
                    error_msg = str(e) if e and str(e) != "None" and str(e).strip() else "Unknown error occurred during file processing"
                    st.error(f"❌ Error processing files: {error_msg}")
                    st.info("💡 Please ensure you've uploaded valid GSC Queries Compare file")
                    # Show column preview for debugging
                    try:
                        if 'gsc_df' in locals() and gsc_df is not None:
                            with st.expander("🔍 Debug: Detected Columns"):
                                st.write(f"Columns found: {', '.join(list(gsc_df.columns)[:15])}")
                    except:
                        pass
                    # Show detailed error in expander
                    with st.expander("🔍 Show Detailed Error"):
                        import traceback
                        st.code(traceback.format_exc(), language='python')
    else:
        st.info("📤 Please upload a GSC Queries Compare file to begin analysis")

def validate_query_data(df):
    """Validate the GSC Queries Compare data"""
    
    # Find query column
    query_col = find_column(df.columns, ['top queries', 'query', 'queries'])
    
    # Find clicks columns (current and previous)
    clicks_current = find_column(df.columns, ['last 3 months clicks', 'clicks']) or \
                    find_column(df.columns, ['click']) and find_column(df.columns, ['last 3', 'current', 'now'])
    
    clicks_previous = find_column(df.columns, ['previous 3 months clicks', 'same period last year clicks']) or \
                     find_column(df.columns, ['click']) and find_column(df.columns, ['previous', 'prev', 'last year'])
    
    missing_columns = []
    if not query_col:
        missing_columns.append('Query/Top Queries')
    if not clicks_current:
        missing_columns.append('Current Period Clicks')
    if not clicks_previous:
        missing_columns.append('Previous Period Clicks')
    
    if missing_columns:
        return False, f"❌ Missing required columns: {missing_columns}. Available columns: {list(df.columns)[:10]}"
    
    if len(df) == 0:
        return False, "❌ File appears to be empty"
    
    return True, "✅ Data validation passed"

def analyze_query_performance(gsc_df, semrush_df=None):
    """Analyze query performance following the prototype methodology"""
    
    # Find columns with flexible matching (handles normalized column names)
    def pick_column(columns, must_include=None, any_of=None):
        must_include = [t.lower().replace(" ", "_") for t in (must_include or [])]
        any_of = [t.lower().replace(" ", "_") for t in (any_of or [])]
        for c in columns:
            # Normalize column name for comparison (handle both original and normalized)
            lc = str(c).lower().replace(" ", "_").replace("-", "_")
            if all(t in lc for t in must_include) and (not any_of or any(t in lc for t in any_of)):
                return c
        return None
    
    cols = list(gsc_df.columns)
    
    # Find GSC columns - try multiple patterns to handle various naming conventions
    query_col = find_column(cols, ['top_queries', 'top queries', 'query', 'queries'])
    
    # Try various patterns for clicks columns (handles normalized names like last_3_months_clicks)
    clicks_now = (find_column(cols, ['last_3_months_clicks', 'last 3 months clicks', 'last_3_months', 'current_clicks', 'clicks_now', 'clicks_current']) or
                  pick_column(cols, must_include=['click'], any_of=['last_3', 'last3', 'current', 'now']) or
                  pick_column(cols, must_include=['last', '3', 'month'], any_of=['click']))
    
    clicks_prev = (find_column(cols, ['same_period_last_year_clicks', 'same period last year clicks', 'previous_3_months_clicks', 'previous 3 months clicks', 'clicks_prev', 'clicks_previous']) or
                   pick_column(cols, must_include=['click'], any_of=['previous', 'prev', 'same_period', 'sameperiod', 'last_year']) or
                   pick_column(cols, must_include=['previous', '3', 'month'], any_of=['click']))
    
    # Try various patterns for impressions columns
    impr_now = (find_column(cols, ['last_3_months_impressions', 'last 3 months impressions', 'impressions_now', 'impressions_current']) or
                pick_column(cols, must_include=['impression'], any_of=['last_3', 'last3', 'current', 'now']) or
                pick_column(cols, must_include=['last', '3', 'month'], any_of=['impression']))
    
    impr_prev = (find_column(cols, ['same_period_last_year_impressions', 'same period last year impressions', 'previous_3_months_impressions', 'previous 3 months impressions', 'impressions_prev']) or
                 pick_column(cols, must_include=['impression'], any_of=['previous', 'prev', 'same_period', 'sameperiod', 'last_year']) or
                 pick_column(cols, must_include=['previous', '3', 'month'], any_of=['impression']))
    
    # Validate required columns
    missing_cols = []
    if not query_col:
        missing_cols.append('Query/Top Queries')
    if not clicks_now:
        missing_cols.append('Current Period Clicks (Last 3 months clicks)')
    if not clicks_prev:
        missing_cols.append('Previous Period Clicks (Previous 3 months clicks or Same period last year clicks)')
    
    if missing_cols:
        available = ", ".join(cols[:15])
        raise ValueError(f"Missing required columns: {missing_cols}. Available columns: {available}...")
    
    # Build working dataframe
    work_df = pd.DataFrame()
    work_df['Query'] = gsc_df[query_col].astype(str).str.strip()
    work_df['Clicks_Now'] = pd.to_numeric(gsc_df[clicks_now], errors='coerce')
    work_df['Clicks_Prev'] = pd.to_numeric(gsc_df[clicks_prev], errors='coerce')
    
    if impr_now and impr_prev:
        work_df['Impr_Now'] = pd.to_numeric(gsc_df[impr_now], errors='coerce')
        work_df['Impr_Prev'] = pd.to_numeric(gsc_df[impr_prev], errors='coerce')
    else:
        work_df['Impr_Now'] = np.nan
        work_df['Impr_Prev'] = np.nan
    
    # Aggregate by query (in case of duplicates)
    agg_df = work_df.groupby('Query', as_index=False).agg({
        'Clicks_Now': 'sum',
        'Clicks_Prev': 'sum', 
        'Impr_Now': 'sum',
        'Impr_Prev': 'sum'
    })
    
    # Calculate deltas and percentages
    agg_df['Clicks_Delta'] = (agg_df['Clicks_Now'] - agg_df['Clicks_Prev']).round(2)
    agg_df['Impr_Delta'] = (agg_df['Impr_Now'] - agg_df['Impr_Prev']).round(2)
    agg_df['Clicks_Pct_Change'] = np.where(agg_df['Clicks_Prev'] > 0, 
                                          (agg_df['Clicks_Delta'] / agg_df['Clicks_Prev'] * 100).round(2), 
                                          np.nan)
    agg_df['Impr_Pct_Change'] = np.where(agg_df['Impr_Prev'] > 0,
                                        (agg_df['Impr_Delta'] / agg_df['Impr_Prev'] * 100).round(2),
                                        np.nan)
    
    # Calculate CTR for both periods
    agg_df['CTR_Now'] = np.where(agg_df['Impr_Now'] > 0, 
                                (agg_df['Clicks_Now'] / agg_df['Impr_Now'] * 100).round(2), 
                                np.nan)
    agg_df['CTR_Prev'] = np.where(agg_df['Impr_Prev'] > 0,
                                 (agg_df['Clicks_Prev'] / agg_df['Impr_Prev'] * 100).round(2),
                                 np.nan)
    agg_df['CTR_Delta_PP'] = (agg_df['CTR_Now'] - agg_df['CTR_Prev']).round(2)
    
    # Semrush enrichment if provided
    if semrush_df is not None:
        try:
            kw_col = find_column(semrush_df.columns, ['keyword'])
            pos_col = find_column(semrush_df.columns, ['position'])
            
            if kw_col and pos_col:
                semrush_work = semrush_df.rename(columns={kw_col: 'Keyword', pos_col: 'Position'})
                semrush_work['Query_Lower'] = semrush_work['Keyword'].astype(str).str.lower().str.strip()
                semrush_positions = semrush_work.set_index('Query_Lower')['Position']
                agg_df['Semrush_Position'] = agg_df['Query'].str.lower().map(semrush_positions)
        except:
            pass  # Skip enrichment if there's any issue
    
    # Top winners and losers
    top_winners = agg_df.sort_values('Clicks_Delta', ascending=False).head(20)
    top_losers = agg_df.sort_values('Clicks_Delta', ascending=True).head(20)
    
    # CTR pressure analysis (impressions up, clicks down)
    ctr_pressure = agg_df[(agg_df['Impr_Delta'] > 0) & (agg_df['Clicks_Delta'] < 0)].sort_values(['Impr_Delta', 'Clicks_Delta'], ascending=[False, True])
    
    # CTR wins (clicks up with impressions flat ±5%)
    impr_flat_mask = (agg_df['Impr_Prev'] > 0) & (abs(agg_df['Impr_Delta'] / agg_df['Impr_Prev']) <= 0.05)
    ctr_wins = agg_df[impr_flat_mask & (agg_df['Clicks_Delta'] > 0)].sort_values('Clicks_Delta', ascending=False)
    
    # Summary metrics
    total_clicks_now = agg_df['Clicks_Now'].sum()
    total_clicks_prev = agg_df['Clicks_Prev'].sum()
    total_impr_now = agg_df['Impr_Now'].sum() 
    total_impr_prev = agg_df['Impr_Prev'].sum()
    
    # Calculate CTR delta
    ctr_now = (total_clicks_now / total_impr_now * 100) if total_impr_now > 0 else np.nan
    ctr_prev = (total_clicks_prev / total_impr_prev * 100) if total_impr_prev > 0 else np.nan
    ctr_delta_pp = (ctr_now - ctr_prev) if not pd.isna(ctr_now) and not pd.isna(ctr_prev) else np.nan
    
    return {
        'total_queries': len(agg_df),
        'total_clicks_now': total_clicks_now,  # Add current total for display
        'total_clicks_prev': total_clicks_prev,  # Add previous total
        'total_clicks_delta': total_clicks_now - total_clicks_prev,
        'total_clicks_pct_change': ((total_clicks_now - total_clicks_prev) / total_clicks_prev * 100) if total_clicks_prev > 0 else 0,
        'total_impr_now': total_impr_now,  # Add current impressions
        'total_impr_prev': total_impr_prev,  # Add previous impressions
        'total_impr_delta': total_impr_now - total_impr_prev,
        'total_impr_pct_change': ((total_impr_now - total_impr_prev) / total_impr_prev * 100) if total_impr_prev > 0 else 0,
        'weighted_ctr_now': ctr_now,
        'weighted_ctr_prev': ctr_prev,
        'ctr_delta_pp': ctr_delta_pp,  # Add CTR delta in percentage points
        'top_winners': top_winners,
        'top_losers': top_losers, 
        'ctr_pressure': ctr_pressure.head(25),
        'ctr_wins': ctr_wins.head(25),
        'raw_data': agg_df,
        'has_semrush': 'Semrush_Position' in agg_df.columns
    }

def display_query_results(results):
    """Display query performance analysis results"""
    
    # Key metrics
    st.markdown('<div class="section-header">📈 Query Performance Summary</div>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            label="Total Queries",
            value=f"{results['total_queries']:,}"
        )
    
    with col2:
        # For clicks: positive change is good (green), negative is bad (red)
        # Streamlit's "normal" mode handles this correctly: positive = green, negative = red
        st.metric(
            label="Total Clicks Change",
            value=f"{results['total_clicks_delta']:,}",
            delta=f"{results['total_clicks_pct_change']:+.1f}%",
            delta_color="normal"  # Always normal: positive = green (good), negative = red (bad)
        )
    
    with col3:
        # For impressions: positive change is good (green), negative is bad (red)
        # Streamlit's "normal" mode handles this correctly: positive = green, negative = red
        st.metric(
            label="Total Impressions Change", 
            value=f"{results['total_impr_delta']:,}",
            delta=f"{results['total_impr_pct_change']:+.1f}%",
            delta_color="normal"  # Always normal: positive = green (good), negative = red (bad)
        )
    
    with col4:
        if not pd.isna(results['weighted_ctr_now']) and not pd.isna(results['weighted_ctr_prev']):
            ctr_delta = results['weighted_ctr_now'] - results['weighted_ctr_prev']
            st.metric(
                label="Weighted CTR",
                value=f"{results['weighted_ctr_now']:.2f}%",
                delta=f"{ctr_delta:+.2f}pp",
                help="Site-wide click-through rate (total clicks / total impressions)"
            )
        else:
            st.metric(label="Weighted CTR", value="N/A", help="Impression data not available")
    
    # Top Winners and Losers Charts
    st.markdown('<div class="section-header">📊 Top Query Movers</div>', unsafe_allow_html=True)
    
    # Top Winners Chart
    if not results['top_winners'].empty:
        top_winners_chart = results['top_winners'].head(15).copy()
        
        fig_winners = go.Figure(data=[
            go.Bar(
                y=top_winners_chart['Query'],
                x=top_winners_chart['Clicks_Delta'],
                orientation='h',
                marker_color='#2ecc71',
                text=[f"+{val:,.0f}" for val in top_winners_chart['Clicks_Delta']],
                textposition='outside'
            )
        ])
        
        fig_winners.update_layout(
            title=dict(text='Top 15 Winning Queries by Clicks Δ', font=dict(size=20)),
            xaxis_title='Clicks Δ (Additional Clicks)',
            yaxis_title='Query',
            height=600,
            margin=dict(l=250, r=60, t=80, b=60),
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            yaxis=dict(autorange='reversed')
        )
        
        st.plotly_chart(fig_winners, use_container_width=True, config={'displayModeBar': False})
    
    # Top Losers Chart  
    if not results['top_losers'].empty:
        top_losers_chart = results['top_losers'].head(15).copy()
        
        fig_losers = go.Figure(data=[
            go.Bar(
                y=top_losers_chart['Query'],
                x=top_losers_chart['Clicks_Delta'],
                orientation='h', 
                marker_color='#e74c3c',
                text=[f"{val:,.0f}" for val in top_losers_chart['Clicks_Delta']],
                textposition='outside'
            )
        ])
        
        fig_losers.update_layout(
            title=dict(text='Top 15 Losing Queries by Clicks Δ', font=dict(size=20)),
            xaxis_title='Clicks Δ (Lost Clicks)',
            yaxis_title='Query',
            height=600,
            margin=dict(l=250, r=60, t=80, b=60),
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            yaxis=dict(autorange='reversed')
        )
        
        st.plotly_chart(fig_losers, use_container_width=True, config={'displayModeBar': False})
    
    # CTR Analysis Tables
    analysis_tab1, analysis_tab2 = st.tabs(["🎯 CTR Pressure Analysis", "🏆 CTR/Ranking Wins"])
    
    with analysis_tab1:
        st.markdown("### CTR/SERP Headwinds - Impressions ↑ but Clicks ↓")
        st.markdown("*Visibility rose but clicks fell - likely CTR or SERP-feature pressure*")
        
        if not results['ctr_pressure'].empty:
            display_cols = ['Query', 'Clicks_Prev', 'Clicks_Now', 'Clicks_Delta', 'Impr_Prev', 'Impr_Now', 'Impr_Delta', 'CTR_Prev', 'CTR_Now', 'CTR_Delta_PP']
            if results['has_semrush']:
                display_cols.append('Semrush_Position')
            
            available_cols = [col for col in display_cols if col in results['ctr_pressure'].columns]
            st.dataframe(results['ctr_pressure'][available_cols], use_container_width=True, hide_index=True, height=400)
        else:
            st.info("No queries found with impression gains but click losses")
    
    with analysis_tab2:
        st.markdown("### Likely CTR/Ranking Wins - Clicks ↑ with Impressions ~Flat (±5%)")
        st.markdown("*Clicks rose without much impression change - better CTR or rankings*")
        
        if not results['ctr_wins'].empty:
            display_cols = ['Query', 'Clicks_Prev', 'Clicks_Now', 'Clicks_Delta', 'Impr_Prev', 'Impr_Now', 'Impr_Delta', 'CTR_Prev', 'CTR_Now', 'CTR_Delta_PP']
            if results['has_semrush']:
                display_cols.append('Semrush_Position')
            
            available_cols = [col for col in display_cols if col in results['ctr_wins'].columns]
            st.dataframe(results['ctr_wins'][available_cols], use_container_width=True, hide_index=True, height=400)
        else:
            st.info("No queries found with click gains and flat impressions")
    
    # Strategic insights - AI-powered
    st.markdown('<div class="section-header">💡 Strategic Insights</div>', unsafe_allow_html=True)
    
    st.info("💡 **AI-Powered Insights:** Strategic insights are generated using AI. Set your OpenAI API key as an environment variable `OPENAI_API_KEY` or in Streamlit secrets.")
    
    with st.spinner("🤖 Generating AI-powered strategic insights..."):
        analysis_summary = create_query_analysis_summary(results)
        insights = generate_chatgpt_insights(analysis_summary, "queries")
        st.markdown(f'<div class="insight-box">{insights.replace(chr(10), "<br>")}</div>', unsafe_allow_html=True)
    
    # Download section
    st.markdown('<div class="section-header">📥 Download Results</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        summary_report = create_query_summary_report(results)
        st.download_button(
            label="📄 Download Query Analysis Report",
            data=summary_report,
            file_name=f"query_performance_analysis_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain"
        )
    
    with col2:
        # Convert top movers to CSV
        csv_buffer = io.StringIO()
        combined_queries = pd.concat([
            results['top_winners'].head(15).assign(Type='Winner'),
            results['top_losers'].head(15).assign(Type='Loser')
        ])
        combined_queries.to_csv(csv_buffer, index=False)
        
        st.download_button(
            label="📊 Download Top Queries (CSV)",
            data=csv_buffer.getvalue(),
            file_name=f"query_top_movers_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv"
        )

def generate_query_insights(results):
    """Generate strategic insights from query performance analysis"""
    insights = []
    
    clicks_delta = results['total_clicks_delta']
    impr_delta = results['total_impr_delta']
    clicks_pct = results['total_clicks_pct_change']
    impr_pct = results['total_impr_pct_change']
    
    # Overall trend analysis
    if clicks_delta < 0 and impr_delta > 0:
        insights.append(f"<b>🔴 CTR Pressure Pattern:</b> Clicks down {abs(clicks_delta):,} ({clicks_pct:.1f}%) while impressions up {impr_delta:,} ({impr_pct:.1f}%) - likely SERP feature or competitive pressure rather than demand drop.")
    elif clicks_delta < 0 and impr_delta < 0:
        insights.append(f"<b>🟡 Demand Decline:</b> Both clicks ({clicks_pct:.1f}%) and impressions ({impr_pct:.1f}%) declined - check seasonality, page losses, or indexing issues.")
    elif clicks_delta > 0 and impr_delta > 0:
        insights.append(f"<b>🟢 Broad Growth:</b> Both clicks (+{clicks_delta:,}) and impressions (+{impr_delta:,}) increased - strong overall SEO momentum.")
    elif clicks_delta > 0 and impr_delta <= 0:
        insights.append(f"<b>🟢 Efficiency Gains:</b> Clicks up {clicks_delta:,} despite flat/declining impressions - improved CTR or rankings.")
    
    # CTR analysis
    if not pd.isna(results['weighted_ctr_now']) and not pd.isna(results['weighted_ctr_prev']):
        ctr_change = results['weighted_ctr_now'] - results['weighted_ctr_prev']
        if abs(ctr_change) > 0.5:
            direction = "improved" if ctr_change > 0 else "declined"
            insights.append(f"<b>📊 CTR Trend:</b> Site-wide CTR {direction} by {abs(ctr_change):.2f} percentage points - {('strong positive signal' if ctr_change > 0 else 'needs investigation')}.")
    
    # Pressure vs wins analysis
    pressure_count = len(results['ctr_pressure'])
    wins_count = len(results['ctr_wins'])
    
    if pressure_count > wins_count * 2:
        insights.append(f"<b>⚠️ SERP Pressure:</b> {pressure_count} queries show CTR pressure vs {wins_count} clear wins - focus on snippet optimization and competitive analysis.")
    elif wins_count > pressure_count:
        insights.append(f"<b>🎯 Ranking Success:</b> {wins_count} queries show clear CTR/ranking wins vs {pressure_count} under pressure - momentum is positive.")
    
    # Semrush context
    if results['has_semrush']:
        insights.append("<b>🔍 Position Context:</b> Semrush position data included for additional ranking context on query changes.")
    
    return "<br><br>".join(insights)

def create_query_summary_report(results):
    """Create downloadable query analysis report"""
    
    ctr_change = results['weighted_ctr_now'] - results['weighted_ctr_prev'] if (not pd.isna(results['weighted_ctr_now']) and not pd.isna(results['weighted_ctr_prev'])) else np.nan
    
    report = f"""
QUERY PERFORMANCE ANALYSIS REPORT
Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

===========================================
EXECUTIVE SUMMARY
===========================================

Total Queries Analyzed: {results['total_queries']:,}

Performance Changes:
• Total Clicks Change: {results['total_clicks_delta']:,} ({results['total_clicks_pct_change']:+.1f}%)
• Total Impressions Change: {results['total_impr_delta']:,} ({results['total_impr_pct_change']:+.1f}%)
• Weighted CTR Change: {ctr_change:+.2f}pp (if available)

Analysis Segments:
• CTR Pressure Cases: {len(results['ctr_pressure'])} queries
• CTR/Ranking Wins: {len(results['ctr_wins'])} queries

===========================================
TOP WINNING QUERIES (Sample)
===========================================

"""
    
    for _, row in results['top_winners'].head(10).iterrows():
        semrush_info = f" | Pos: {row['Semrush_Position']}" if 'Semrush_Position' in row and pd.notna(row['Semrush_Position']) else ""
        report += f"• {row['Query']} | +{row['Clicks_Delta']:.0f} clicks ({row['Clicks_Pct_Change']:+.1f}%){semrush_info}\n"
    
    report += f"""

===========================================
TOP LOSING QUERIES (Sample)
===========================================

"""
    
    for _, row in results['top_losers'].head(10).iterrows():
        semrush_info = f" | Pos: {row['Semrush_Position']}" if 'Semrush_Position' in row and pd.notna(row['Semrush_Position']) else ""
        report += f"• {row['Query']} | {row['Clicks_Delta']:.0f} clicks ({row['Clicks_Pct_Change']:+.1f}%){semrush_info}\n"
    
    report += f"""

===========================================
STRATEGIC INSIGHTS
===========================================

{generate_query_insights(results).replace('<b>', '').replace('</b>', '').replace('<br><br>', '\n\n').replace('🟢', '• ').replace('🟡', '• ').replace('🔴', '• ').replace('🎯', '• ').replace('⚠️', '• ').replace('📊', '• ').replace('🔍', '• ')}

===========================================
"""
    
    return report

def infer_competitor_domain(filename, df):
    """Infer competitor domain from filename or data"""
    
    # Try to extract domain from filename
    import re
    domain_match = re.search(r'([a-z0-9.-]+\.[a-z]{2,})', filename.lower())
    if domain_match:
        domain = domain_match.group(1).replace("www.", "")
        return domain
    
    # Try to extract from URL column if present
    url_col = find_column(df.columns, ['url', 'page', 'landing page'])
    if url_col and not df[url_col].empty:
        try:
            from urllib.parse import urlparse
            sample_url = df[url_col].dropna().iloc[0]
            parsed = urlparse(str(sample_url))
            domain = parsed.netloc.lower().replace("www.", "")
            return domain if domain else None
        except:
            pass
    
    return None

def validate_competitor_data(competitors_df, positions_df):
    """Validate competitor analysis data"""
    
    # Check competitors file
    domain_col = find_column(competitors_df.columns, ['domain', 'competitor', 'competitor domain'])
    
    if not domain_col:
        return False, f"❌ Competitors file missing domain column. Available: {list(competitors_df.columns)[:10]}"
    
    # Check positions file
    keyword_col = find_column(positions_df.columns, ['keyword'])
    position_col = find_column(positions_df.columns, ['position'])
    
    missing = []
    if not keyword_col:
        missing.append('Keyword')
    if not position_col:
        missing.append('Position')
    
    if missing:
        return False, f"❌ Your Positions file missing: {missing}. Available: {list(positions_df.columns)[:10]}"
    
    if len(competitors_df) == 0 or len(positions_df) == 0:
        return False, "❌ One or both files appear to be empty"
    
    return True, "✅ Data validation passed"

def analyze_competitor_gaps(competitors_df, your_positions_df, competitor_data):
    """Analyze competitor gaps following the prototype methodology"""
    
    # 1. Process competitors data
    domain_col = find_column(competitors_df.columns, ['domain', 'competitor', 'competitor domain'])
    relevance_col = find_column(competitors_df.columns, ['competitor relevance', 'relevance'])
    common_kw_col = find_column(competitors_df.columns, ['common keywords', 'common kws', 'common'])
    traffic_col = find_column(competitors_df.columns, ['organic keywords', 'organic traffic'])
    
    competitors_work = pd.DataFrame()
    competitors_work['Domain'] = competitors_df[domain_col].astype(str).str.strip()
    
    if relevance_col:
        competitors_work['Relevance'] = pd.to_numeric(competitors_df[relevance_col], errors='coerce')
    if common_kw_col:
        competitors_work['Common_Keywords'] = pd.to_numeric(competitors_df[common_kw_col], errors='coerce')
    if traffic_col:
        competitors_work['Organic_Traffic'] = pd.to_numeric(competitors_df[traffic_col], errors='coerce')
    
    # Sort by relevance or common keywords
    if 'Relevance' in competitors_work.columns and competitors_work['Relevance'].notna().any():
        competitors_work = competitors_work.sort_values('Relevance', ascending=False)
    elif 'Common_Keywords' in competitors_work.columns:
        competitors_work = competitors_work.sort_values('Common_Keywords', ascending=False)
    
    # Get top 5 competitors
    top_competitors = competitors_work.head(5)
    
    # 2. Process your positions
    keyword_col = find_column(your_positions_df.columns, ['keyword'])
    position_col = find_column(your_positions_df.columns, ['position'])
    url_col = find_column(your_positions_df.columns, ['url', 'page'])
    
    your_positions_work = pd.DataFrame()
    your_positions_work['Keyword'] = your_positions_df[keyword_col].astype(str).str.strip().str.lower()
    your_positions_work['Your_Position'] = pd.to_numeric(your_positions_df[position_col], errors='coerce')
    
    if url_col:
        your_positions_work['Your_URL'] = your_positions_df[url_col].astype(str)
    
    # Keep best position per keyword
    your_positions_work = (your_positions_work.groupby('Keyword')
                          .agg({'Your_Position': 'min', 'Your_URL': 'first' if url_col else 'count'})
                          .reset_index())
    
    # 3. Build gap analysis for competitors with position data
    gap_analysis = []
    summary_counts = []
    
    for _, competitor in top_competitors.iterrows():
        domain = competitor['Domain']
        
        # Check if we have position data for this competitor
        if domain in competitor_data:
            comp_df = competitor_data[domain]
            
            # Process competitor positions
            comp_kw_col = find_column(comp_df.columns, ['keyword'])
            comp_pos_col = find_column(comp_df.columns, ['position'])
            
            if comp_kw_col and comp_pos_col:
                comp_positions = pd.DataFrame()
                comp_positions['Keyword'] = comp_df[comp_kw_col].astype(str).str.strip().str.lower()
                comp_positions['Comp_Position'] = pd.to_numeric(comp_df[comp_pos_col], errors='coerce')
                
                # Keep best position per keyword
                comp_positions = (comp_positions.groupby('Keyword')['Comp_Position'].min().reset_index())
                
                # Join with your positions
                gap_df = your_positions_work.merge(comp_positions, on='Keyword', how='inner')
                
                # Calculate effective ranks (treat 0/NaN as worst)
                def effective_rank(series):
                    return np.where((series <= 0) | pd.isna(series), 1000, series)
                
                your_eff = effective_rank(gap_df['Your_Position'])
                comp_eff = effective_rank(gap_df['Comp_Position'])
                
                gap_df['Comp_Outranks'] = (comp_eff < your_eff)
                gap_df['Gap'] = (your_eff - comp_eff).astype(float)
                gap_df['Competitor'] = domain
                
                gap_analysis.append(gap_df)
                
                # Summary counts
                outranked = int(gap_df['Comp_Outranks'].sum())
                we_outrank = int((your_eff < comp_eff).sum())
                ties = int((your_eff == comp_eff).sum())
                total_matched = len(gap_df)
                
                summary_counts.append({
                    'Competitor': domain,
                    'Outranked_Us': outranked,
                    'We_Outrank': we_outrank,
                    'Ties': ties,
                    'Matched_Keywords': total_matched,
                    'Relevance': competitor.get('Relevance', np.nan),
                    'Common_Keywords': competitor.get('Common_Keywords', np.nan)
                })
    
    # Combine all gap data
    all_gaps = pd.concat(gap_analysis) if gap_analysis else pd.DataFrame()
    summary_df = pd.DataFrame(summary_counts)
    
    # 4. Focus on losing queries if available (check if query analysis was run)
    losing_focus = pd.DataFrame()
    if hasattr(st.session_state, 'query_results') and 'top_losers' in st.session_state.query_results:
        try:
            losing_queries = st.session_state.query_results['top_losers']['Query'].str.lower().tolist()[:50]
            if not all_gaps.empty:
                losing_focus = all_gaps[all_gaps['Keyword'].isin(losing_queries) & all_gaps['Comp_Outranks']]
        except:
            pass  # Skip if query data not available or incompatible
    
    return {
        'top_competitors': top_competitors,
        'your_keywords_count': len(your_positions_work),
        'gap_analysis': all_gaps,
        'summary_counts': summary_df,
        'losing_query_focus': losing_focus,
        'has_position_data': len(gap_analysis) > 0,
        'competitors_with_data': list(competitor_data.keys())
    }

def display_competitor_results(results):
    """Display competitor analysis results"""
    
    # Key metrics
    st.markdown('<div class="section-header">🏆 Competitive Landscape Summary</div>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            label="Top Competitors",
            value=len(results['top_competitors']),
            help="Based on relevance and keyword overlap"
        )
    
    with col2:
        st.metric(
            label="Your Keywords",
            value=f"{results['your_keywords_count']:,}",
            help="Total keywords in your position data"
        )
    
    with col3:
        if results['has_position_data']:
            total_gaps = len(results['gap_analysis'])
            st.metric(
                label="Gap Analysis Keywords",
                value=f"{total_gaps:,}",
                help="Keywords with competitive comparison data"
            )
        else:
            st.metric(
                label="Gap Analysis",
                value="Limited",
                help="Upload competitor position files for detailed gaps"
            )
    
    with col4:
        competitors_with_data = len(results['competitors_with_data'])
        st.metric(
            label="Competitors with Data",
            value=competitors_with_data,
            help="Competitors with uploaded position data"
        )
    
    # Top Competitors Table
    st.markdown('<div class="section-header">🎯 Top Search Competitors</div>', unsafe_allow_html=True)
    st.markdown("*Domains most similar to you in search results (not necessarily business competitors)*")
    
    display_cols = ['Domain']
    if 'Relevance' in results['top_competitors'].columns:
        display_cols.append('Relevance')
    if 'Common_Keywords' in results['top_competitors'].columns:
        display_cols.append('Common_Keywords')
    if 'Organic_Traffic' in results['top_competitors'].columns:
        display_cols.append('Organic_Traffic')
    
    competitors_display = results['top_competitors'][display_cols].copy()
    st.dataframe(competitors_display, use_container_width=True, hide_index=True)
    
    # Gap Analysis Results
    if results['has_position_data']:
        st.markdown('<div class="section-header">📊 Competitive Gap Analysis</div>', unsafe_allow_html=True)
        
        # Summary counts
        if not results['summary_counts'].empty:
            st.markdown("**Where competitors outrank you:**")
            
            summary_display = results['summary_counts'].copy()
            if 'Relevance' in summary_display.columns:
                summary_display = summary_display.drop('Relevance', axis=1)
            if 'Common_Keywords' in summary_display.columns:
                summary_display = summary_display.drop('Common_Keywords', axis=1)
                
            st.dataframe(summary_display, use_container_width=True, hide_index=True)
            
            # Outrank chart
            if len(results['summary_counts']) > 1:
                fig_outrank = go.Figure(data=[
                    go.Bar(
                        y=results['summary_counts']['Competitor'],
                        x=results['summary_counts']['Outranked_Us'],
                        orientation='h',
                        marker_color='#e74c3c',
                        text=results['summary_counts']['Outranked_Us'],
                        textposition='outside'
                    )
                ])
                
                fig_outrank.update_layout(
                    title=dict(text='Keywords Where Competitor Outranks You', font=dict(size=20)),
                    xaxis_title='Keywords Count',
                    yaxis_title='Competitor',
                    height=400,
                    margin=dict(l=150, r=60, t=80, b=60),
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)',
                    yaxis=dict(autorange='reversed')
                )
                
                st.plotly_chart(fig_outrank, use_container_width=True, config={'displayModeBar': False})
        
        # Detailed gap opportunities
        if not results['gap_analysis'].empty:
            st.markdown("### 🎯 Specific Gap Opportunities")
            st.markdown("*Keywords where competitors rank better than you - sorted by largest gaps*")
            
            gap_opportunities = results['gap_analysis'][results['gap_analysis']['Comp_Outranks']].copy()
            gap_opportunities = gap_opportunities.sort_values(['Competitor', 'Gap'], ascending=[True, False])
            
            display_cols = ['Competitor', 'Keyword', 'Your_Position', 'Comp_Position', 'Gap']
            if 'Your_URL' in gap_opportunities.columns:
                display_cols.append('Your_URL')
            
            gap_display = gap_opportunities[display_cols].head(50)
            gap_display.columns = ['Competitor', 'Keyword', 'Your Position', 'Competitor Position', 'Gap'] + (['Your URL'] if len(display_cols) > 5 else [])
            
            st.dataframe(gap_display, use_container_width=True, hide_index=True, height=400)
        
        # Focus on losing queries
        if not results['losing_query_focus'].empty:
            st.markdown("### 📉 Competitive Pressure on Declining Queries")
            st.markdown("*Competitors outranking you on queries that lost traffic*")
            
            losing_display = results['losing_query_focus'][['Competitor', 'Keyword', 'Your_Position', 'Comp_Position', 'Gap']].copy()
            losing_display.columns = ['Competitor', 'Keyword', 'Your Position', 'Competitor Position', 'Gap']
            
            st.dataframe(losing_display, use_container_width=True, hide_index=True, height=300)
        elif results['has_position_data']:
            st.info("💡 Run Query Analysis first to see competitive pressure on your declining queries")
    
    else:
        st.markdown('<div class="section-header">📊 Competitive Intelligence Analysis</div>', unsafe_allow_html=True)
        st.markdown("*Analysis based on competitor overlap and your current positions*")
        
        # Enhanced analysis without detailed position data
        enhanced_results = analyze_competitive_intelligence(results)
        display_enhanced_competitor_analysis(enhanced_results)

def analyze_competitive_intelligence(results):
    """Enhanced analysis using just competitors and your positions data"""
    
    # Get your top keywords by position quality
    your_positions = results['gap_analysis'] if hasattr(results, 'gap_analysis') and not results['gap_analysis'].empty else pd.DataFrame()
    
    # If we don't have position data in results, we need to get it differently
    # Let's enhance the basic competitor results
    competitors = results['top_competitors']
    
    # 1. Market positioning analysis
    positioning_analysis = analyze_market_positioning(competitors)
    
    # 2. Competitive pressure estimation
    pressure_analysis = estimate_competitive_pressure(competitors, results['your_keywords_count'])
    
    # 3. Opportunity sizing
    opportunity_analysis = analyze_opportunity_sizing(competitors)
    
    return {
        **results,
        'market_positioning': positioning_analysis,
        'competitive_pressure': pressure_analysis,
        'opportunity_analysis': opportunity_analysis
    }

def analyze_market_positioning(competitors_df):
    """Analyze market positioning based on competitor metrics"""
    
    if competitors_df.empty:
        return {}
    
    # Calculate market share indicators
    total_common_kw = competitors_df['Common_Keywords'].sum() if 'Common_Keywords' in competitors_df.columns else 0
    
    positioning = []
    for _, comp in competitors_df.iterrows():
        common_kw = comp.get('Common_Keywords', 0)
        relevance = comp.get('Relevance', 0)
        
        # Market share estimation
        kw_share = (common_kw / total_common_kw * 100) if total_common_kw > 0 else 0
        
        # Threat level based on relevance and overlap
        if relevance > 80 and common_kw > 1000:
            threat_level = "High"
        elif relevance > 60 and common_kw > 500:
            threat_level = "Medium"
        else:
            threat_level = "Low"
        
        positioning.append({
            'Domain': comp['Domain'],
            'Keyword_Share': kw_share,
            'Threat_Level': threat_level,
            'Common_Keywords': common_kw,
            'Relevance': relevance
        })
    
    return pd.DataFrame(positioning)

def estimate_competitive_pressure(competitors_df, your_keyword_count):
    """Estimate competitive pressure across different areas"""
    
    if competitors_df.empty:
        return {}
    
    # High competition indicators
    high_competition_count = len(competitors_df[
        (competitors_df.get('Relevance', 0) > 70) & 
        (competitors_df.get('Common_Keywords', 0) > your_keyword_count * 0.3)
    ])
    
    # Market saturation estimate
    avg_common_keywords = competitors_df['Common_Keywords'].mean() if 'Common_Keywords' in competitors_df.columns else 0
    saturation_level = "High" if avg_common_keywords > your_keyword_count * 0.5 else "Medium" if avg_common_keywords > your_keyword_count * 0.2 else "Low"
    
    # Competitive intensity
    total_competitors = len(competitors_df)
    high_relevance_competitors = len(competitors_df[competitors_df.get('Relevance', 0) > 60])
    
    intensity_score = (high_relevance_competitors / total_competitors * 100) if total_competitors > 0 else 0
    
    return {
        'high_competition_competitors': high_competition_count,
        'market_saturation': saturation_level,
        'competitive_intensity': intensity_score,
        'avg_keyword_overlap': avg_common_keywords
    }

def analyze_opportunity_sizing(competitors_df):
    """Analyze opportunity sizing based on competitor data"""
    
    if competitors_df.empty:
        return {}
    
    # Opportunity categories
    opportunities = []
    
    for _, comp in competitors_df.iterrows():
        common_kw = comp.get('Common_Keywords', 0)
        relevance = comp.get('Relevance', 0)
        
        # Determine opportunity type
        if relevance < 50 and common_kw > 500:
            opp_type = "Keyword Expansion"
            priority = "High"
        elif relevance > 80 and common_kw < 200:
            opp_type = "Niche Domination"
            priority = "Medium"
        elif relevance > 70 and common_kw > 1000:
            opp_type = "Head-to-Head Competition" 
            priority = "High"
        else:
            opp_type = "Market Monitoring"
            priority = "Low"
        
        opportunities.append({
            'Domain': comp['Domain'],
            'Opportunity_Type': opp_type,
            'Priority': priority,
            'Estimated_Keywords': common_kw,
            'Competitive_Strength': "Strong" if relevance > 70 else "Moderate" if relevance > 40 else "Weak"
        })
    
    return pd.DataFrame(opportunities)

def display_enhanced_competitor_analysis(results):
    """Display enhanced competitor analysis without detailed position data"""
    
    # Market Positioning Analysis
    if 'market_positioning' in results and not results['market_positioning'].empty:
        st.markdown("### 🎯 Market Positioning Analysis")
        
        # Market share chart
        positioning = results['market_positioning']
        
        fig_market = go.Figure(data=[
            go.Bar(
                x=positioning['Domain'],
                y=positioning['Keyword_Share'],
                marker_color=['#e74c3c' if threat == 'High' else '#f39c12' if threat == 'Medium' else '#27ae60' 
                            for threat in positioning['Threat_Level']],
                text=[f"{share:.1f}%" for share in positioning['Keyword_Share']],
                textposition='outside'
            )
        ])
        
        fig_market.update_layout(
            title=dict(text='Estimated Market Share by Keyword Overlap', font=dict(size=18)),
            xaxis_title='Competitor',
            yaxis_title='Keyword Overlap Share (%)',
            height=400,
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            xaxis_tickangle=-45
        )
        
        st.plotly_chart(fig_market, use_container_width=True, config={'displayModeBar': False})
        
        # Positioning table
        st.dataframe(positioning, use_container_width=True, hide_index=True)
    
    # Competitive Pressure Analysis
    if 'competitive_pressure' in results:
        pressure = results['competitive_pressure']
        
        st.markdown("### ⚠️ Competitive Pressure Assessment")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric(
                label="High Competition",
                value=pressure['high_competition_competitors'],
                help="Competitors with high relevance and significant keyword overlap"
            )
        
        with col2:
            color = "🔴" if pressure['market_saturation'] == "High" else "🟡" if pressure['market_saturation'] == "Medium" else "🟢"
            st.metric(
                label="Market Saturation",
                value=f"{color} {pressure['market_saturation']}",
                help="Level of competitive saturation in your market"
            )
        
        with col3:
            st.metric(
                label="Competitive Intensity",
                value=f"{pressure['competitive_intensity']:.0f}%",
                help="Percentage of competitors with high relevance scores"
            )
        
        with col4:
            st.metric(
                label="Avg Keyword Overlap",
                value=f"{pressure['avg_keyword_overlap']:.0f}",
                help="Average keywords shared with competitors"
            )
    
    # Opportunity Analysis
    if 'opportunity_analysis' in results and not results['opportunity_analysis'].empty:
        st.markdown("### 🎯 Strategic Opportunities")
        
        opportunities = results['opportunity_analysis']
        
        # Opportunity type distribution
        opp_counts = opportunities['Opportunity_Type'].value_counts()
        
        fig_opp = go.Figure(data=[
            go.Pie(
                labels=opp_counts.index,
                values=opp_counts.values,
                hole=0.3,
                marker_colors=['#3498db', '#e74c3c', '#f39c12', '#27ae60']
            )
        ])
        
        fig_opp.update_layout(
            title=dict(text='Strategic Opportunity Distribution', font=dict(size=18)),
            height=400,
            annotations=[dict(text='Opportunities', x=0.5, y=0.5, font_size=16, showarrow=False)]
        )
        
        st.plotly_chart(fig_opp, use_container_width=True, config={'displayModeBar': False})
        
        # Opportunities table
        st.markdown("**Detailed Opportunity Analysis:**")
        st.dataframe(opportunities, use_container_width=True, hide_index=True)
    
    # Competitive Intelligence Summary
    st.markdown("### 📋 Competitive Intelligence Summary")
    
    # Generate insights for basic competitor data
    intel_insights = []
    
    if 'market_positioning' in results and not results['market_positioning'].empty:
        top_threat = results['market_positioning'].loc[results['market_positioning']['Keyword_Share'].idxmax()]
        intel_insights.append(f"<b>🎯 Market Leader:</b> {top_threat['Domain']} has the highest keyword overlap ({top_threat['Keyword_Share']:.1f}%) and {top_threat['Threat_Level'].lower()} threat level.")
    
    if 'competitive_pressure' in results:
        pressure = results['competitive_pressure']
        intel_insights.append(f"<b>📊 Market Dynamics:</b> {pressure['market_saturation']} market saturation with {pressure['competitive_intensity']:.0f}% of competitors showing high relevance scores.")
    
    if 'opportunity_analysis' in results and not results['opportunity_analysis'].empty:
        high_priority = len(results['opportunity_analysis'][results['opportunity_analysis']['Priority'] == 'High'])
        intel_insights.append(f"<b>⚡ Action Items:</b> {high_priority} high-priority competitive opportunities identified for immediate strategic focus.")
    
    # Always add this insight
    intel_insights.append("<b>🚀 Next Level:</b> Upload individual competitor position files to unlock detailed keyword gap analysis, specific ranking opportunities, and head-to-head competitive comparisons.")
    
    if intel_insights:
        st.markdown(f'<div class="insight-box">{"<br><br>".join(intel_insights)}</div>', unsafe_allow_html=True)
    
    # Strategic insights - AI-powered
    st.markdown('<div class="section-header">💡 Strategic Insights</div>', unsafe_allow_html=True)
    
    st.info("💡 **AI-Powered Insights:** Strategic insights are generated using AI. Set your OpenAI API key as an environment variable `OPENAI_API_KEY` or in Streamlit secrets.")
    
    with st.spinner("🤖 Generating AI-powered strategic insights..."):
        analysis_summary = create_competitor_analysis_summary(results)
        insights = generate_chatgpt_insights(analysis_summary, "competitors")
        st.markdown(f'<div class="insight-box">{insights.replace(chr(10), "<br>")}</div>', unsafe_allow_html=True)
    
    # Download section
    st.markdown('<div class="section-header">📥 Download Results</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        summary_report = create_competitor_summary_report(results)
        st.download_button(
            label="📄 Download Competitor Analysis Report",
            data=summary_report,
            file_name=f"competitor_analysis_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain"
        )
    
    with col2:
        if results['has_position_data']:
            csv_buffer = io.StringIO()
            results['gap_analysis'].to_csv(csv_buffer, index=False)
            
            st.download_button(
                label="📊 Download Gap Analysis (CSV)",
                data=csv_buffer.getvalue(),
                file_name=f"competitor_gaps_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )
        else:
            st.download_button(
                label="📊 Download Competitors List (CSV)", 
                data=results['top_competitors'].to_csv(index=False),
                file_name=f"top_competitors_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )

def generate_competitor_insights(results):
    """Generate strategic insights from competitor analysis"""
    insights = []
    
    # Competitor identification insights
    competitors_count = len(results['top_competitors'])
    insights.append(f"<b>🎯 Competitive Landscape:</b> Identified {competitors_count} key search competitors based on keyword overlap and relevance.")
    
    # Gap analysis insights
    if results['has_position_data']:
        total_gaps = len(results['gap_analysis'])
        outranked_keywords = len(results['gap_analysis'][results['gap_analysis']['Comp_Outranks']]) if not results['gap_analysis'].empty else 0
        
        if not results['summary_counts'].empty:
            top_threat = results['summary_counts'].loc[results['summary_counts']['Outranked_Us'].idxmax()]
            insights.append(f"<b>🥇 Top Competitive Threat:</b> {top_threat['Competitor']} outranks you on {top_threat['Outranked_Us']} keywords from {top_threat['Matched_Keywords']} matched keywords.")
        
        if outranked_keywords > 0:
            insights.append(f"<b>📊 Gap Opportunities:</b> {outranked_keywords:,} keywords where competitors rank better than you - prime targets for content optimization and competitive analysis.")
        
        # Losing query focus
        if not results['losing_query_focus'].empty:
            losing_count = len(results['losing_query_focus'])
            insights.append(f"<b>🔴 Competitive Pressure:</b> {losing_count} of your declining queries show direct competitive displacement - these require immediate attention.")
        
        # Strategy recommendations
        if outranked_keywords > total_gaps * 0.3:  # If >30% of gaps show we're behind
            insights.append("<b>⚡ Action Priority:</b> High competitive pressure detected. Focus on content depth, technical optimization, and link building for gap keywords.")
    else:
        insights.append("<b>🎯 Next Steps:</b> Upload competitor position files to unlock detailed gap analysis, outrank counts, and specific optimization opportunities.")
    
    return "<br><br>".join(insights)

def create_competitor_summary_report(results):
    """Create downloadable competitor analysis report"""
    
    report = f"""
COMPETITOR GAP ANALYSIS REPORT
Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

===========================================
EXECUTIVE SUMMARY
===========================================

Top Search Competitors Identified: {len(results['top_competitors'])}
Your Keywords Analyzed: {results['your_keywords_count']:,}
Gap Analysis Available: {'Yes' if results['has_position_data'] else 'Limited - upload competitor data'}

===========================================
TOP SEARCH COMPETITORS
===========================================

"""
    
    for _, row in results['top_competitors'].iterrows():
        relevance = f" | Relevance: {row['Relevance']}" if 'Relevance' in row and pd.notna(row['Relevance']) else ""
        common_kw = f" | Common KW: {row['Common_Keywords']:,.0f}" if 'Common_Keywords' in row and pd.notna(row['Common_Keywords']) else ""
        report += f"• {row['Domain']}{relevance}{common_kw}\n"
    
    if results['has_position_data']:
        report += f"""

===========================================
COMPETITIVE OUTRANK ANALYSIS
===========================================

"""
        
        for _, row in results['summary_counts'].iterrows():
            report += f"• {row['Competitor']}: Outranks you on {row['Outranked_Us']} keywords (from {row['Matched_Keywords']} matched)\n"
        
        if not results['gap_analysis'].empty:
            report += f"""

===========================================
TOP GAP OPPORTUNITIES (Sample)
===========================================

"""
            
            gap_sample = results['gap_analysis'][results['gap_analysis']['Comp_Outranks']].head(20)
            for _, row in gap_sample.iterrows():
                report += f"• {row['Keyword']} | You: #{row['Your_Position']} vs {row['Competitor']}: #{row['Comp_Position']} (Gap: {row['Gap']:.0f})\n"
    
    report += f"""

===========================================
STRATEGIC INSIGHTS
===========================================

{generate_competitor_insights(results).replace('<b>', '').replace('</b>', '').replace('<br><br>', '\n\n').replace('🎯', '• ').replace('🥇', '• ').replace('📊', '• ').replace('🔴', '• ').replace('⚡', '• ')}

===========================================
"""
    
    return report

def read_uploaded_file_safe(uploaded_file):
    """Read uploaded CSV or Excel file with better error handling"""
    if uploaded_file is not None:
        file_name = uploaded_file.name.lower()
        try:
            if file_name.endswith('.xlsx') or file_name.endswith('.xls'):
                # Try reading Excel file
                df = pd.read_excel(uploaded_file)
                return df
            else:
                # Try different CSV reading approaches
                try:
                    # First try standard CSV
                    df = pd.read_csv(uploaded_file)
                    return df
                except:
                    # Reset file pointer and try with different separator
                    uploaded_file.seek(0)
                    df = pd.read_csv(uploaded_file, sep=';')
                    return df
        except Exception as e:
            st.error(f"Could not read file {uploaded_file.name}: {str(e)}")
            st.info("Try saving your file as CSV with comma separators, or as Excel format")
            return None
    return None

# REMOVED: traffic_attribution_analysis() and all related helper functions have been removed


def generate_chatgpt_insights(analysis_summary, analysis_type="comprehensive"):
    """Generate strategic insights using ChatGPT API"""
    try:
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            # Try to get from Streamlit secrets
            try:
                api_key = st.secrets.get("OPENAI_API_KEY", "")
            except:
                pass
        
        if not api_key:
            return "⚠️ OpenAI API key not found. Please set OPENAI_API_KEY environment variable or add it to Streamlit secrets to enable AI-powered insights."
        
        client = OpenAI(api_key=api_key)
        
        # Customize prompt based on analysis type
        analysis_context = {
            "visibility": "keyword visibility and ranking distribution analysis",
            "movement": "keyword ranking movement and position changes analysis",
            "pages": "page performance and traffic concentration analysis",
            "queries": "Google Search Console query performance analysis",
            "competitors": "competitive analysis and market positioning",
            "comprehensive": "comprehensive SEO performance analysis"
        }
        
        context = analysis_context.get(analysis_type, "SEO performance analysis")
        
        prompt = f"""You are an expert SEO strategist with 15+ years of experience. Analyze the following {context} and provide strategic, actionable insights that drive business results. Be specific, data-driven, and focus on ROI-impacting recommendations.

{analysis_summary}

Provide a focused analysis with:

1. **Key Findings** (3-4 bullet points)
   - Most important trends and patterns
   - Quantifiable changes with specific numbers
   - What this means for the business

2. **Critical Issues** (3-4 bullet points)
   - Performance problems requiring immediate attention
   - Specific metrics showing decline
   - Root cause hypotheses

3. **Opportunities** (3-4 bullet points)
   - Quick wins and low-hanging fruit
   - Areas with high potential
   - Specific actions to capitalize

4. **Strategic Recommendations** (4-5 actionable items)
   - Prioritized by impact
   - Specific steps to take
   - Expected outcomes and timelines

5. **Priority Actions** (top 3 immediate actions)
   - What to do this week
   - Specific steps
   - Success metrics

Format as clear, professional insights. Use specific numbers from the data. Be actionable and focus on business impact. Use HTML formatting with <b> for emphasis and <br> for line breaks."""
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert SEO strategist providing data-driven insights and recommendations."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1500
        )
        
        return response.choices[0].message.content
    except Exception as e:
        return f"⚠️ Error generating AI insights: {str(e)}. Please check your OpenAI API key configuration."

def create_visibility_analysis_summary(results):
    """Create analysis summary for AI insights - Visibility"""
    total_change = results['total_change']
    total_change_pct = results['total_change_pct']
    top_3_share = results['bucket_changes']['top_3']['current_share']
    top_10_share = results['bucket_changes']['top_3']['current_share'] + results['bucket_changes']['top_4_10']['current_share']
    top_21_plus_share = results['bucket_changes']['top_21_plus']['current_share']
    
    return f"""KEYWORD VISIBILITY ANALYSIS:

Total Keywords: {results['total_current']:,} (current) vs {results['total_previous']:,} (previous)
Net Change: {total_change:+,} keywords ({total_change_pct:+.1f}%)

Ranking Distribution (Current):
- Top 3 Rankings: {results['bucket_changes']['top_3']['current']:,} keywords ({top_3_share:.1f}% of total)
- Top 4-10 Rankings: {results['bucket_changes']['top_4_10']['current']:,} keywords ({results['bucket_changes']['top_4_10']['current_share']:.1f}%)
- Top 11-20 Rankings: {results['bucket_changes']['top_11_20']['current']:,} keywords ({results['bucket_changes']['top_11_20']['current_share']:.1f}%)
- Rankings 21+: {results['bucket_changes']['top_21_plus']['current']:,} keywords ({top_21_plus_share:.1f}%)

Changes by Bucket:
- Top 3: {results['bucket_changes']['top_3']['change']:+,} ({results['bucket_changes']['top_3']['change_pct']:+.1f}%)
- Top 4-10: {results['bucket_changes']['top_4_10']['change']:+,} ({results['bucket_changes']['top_4_10']['change_pct']:+.1f}%)
- Top 11-20: {results['bucket_changes']['top_11_20']['change']:+,} ({results['bucket_changes']['top_11_20']['change_pct']:+.1f}%)
- 21+: {results['bucket_changes']['top_21_plus']['change']:+,} ({results['bucket_changes']['top_21_plus']['change_pct']:+.1f}%)

Top 10 Share: {top_10_share:.1f}% of total keywords
"""

def create_movement_analysis_summary(results):
    """Create analysis summary for AI insights - Movement"""
    improved = results['movement_counts']['improved']
    declined = results['movement_counts']['declined']
    ratio = results['ratio']
    ratio_display = f"{ratio:.2f}" if ratio != np.inf else "∞"
    
    return f"""KEYWORD MOVEMENT ANALYSIS:

Total Keywords Tracked: {results['total_keywords']:,}
Improved Rankings: {improved:,} keywords ({improved/results['total_keywords']*100:.1f}%)
Declined Rankings: {declined:,} keywords ({declined/results['total_keywords']*100:.1f}%)
Improvement Ratio: {ratio_display} (higher is better)
Net Movement: {improved - declined:+,} keywords

Top 3 Sources: {results.get('top3_sources', pd.Series()).to_dict() if not results.get('top3_sources', pd.Series()).empty else 'N/A'}
"""

def create_pages_analysis_summary(results):
    """Create analysis summary for AI insights - Pages"""
    pages_50 = results['pareto_thresholds']['50%']
    pages_80 = results['pareto_thresholds']['80%']
    concentration_50 = (pages_50 / results['total_pages'] * 100) if results['total_pages'] > 0 else 0
    
    return f"""PAGE PERFORMANCE ANALYSIS:

Total Pages Analyzed: {results['total_pages']:,}
Total Estimated Traffic: {results['total_traffic']:,.0f}
Traffic Concentration: {pages_50} pages ({concentration_50:.1f}% of total) drive 50% of traffic
Traffic Concentration: {pages_80} pages ({(pages_80 / results['total_pages'] * 100) if results['total_pages'] > 0 else 0:.1f}% of total) drive 80% of traffic
Long-tail Opportunities: {len(results.get('longtail_opportunities', pd.DataFrame()))} pages identified
"""

def create_query_analysis_summary(results):
    """Create analysis summary for AI insights - Queries"""
    return f"""GSC QUERY PERFORMANCE ANALYSIS:

Total Queries Tracked: {results.get('total_queries', 0):,}
Clicks Change: {results.get('total_clicks_delta', 0):+,} clicks ({results.get('total_clicks_pct_change', 0):+.1f}%)
Impressions Change: {results.get('total_impr_delta', 0):+,} impressions ({results.get('total_impr_pct_change', 0):+.1f}%)
CTR Change: {results.get('ctr_delta_pp', 0):+.2f} percentage points
Top Winners: {len(results.get('top_winners', pd.DataFrame()))} queries with significant gains
Top Losers: {len(results.get('top_losers', pd.DataFrame()))} queries with significant declines
"""

def create_competitor_analysis_summary(results):
    """Create analysis summary for AI insights - Competitors"""
    return f"""COMPETITOR ANALYSIS:

Top Competitors Identified: {len(results.get('top_competitors', []))}
Your Keywords: {results.get('your_keywords_count', 0):,}
Total Outrank Opportunities: {results.get('total_gaps', 0):,}
Competitors with Data: {len(results.get('competitors_with_data', []))}
"""

# REMOVED: comprehensive_report_tab() has been removed


if __name__ == "__main__":
    main()
